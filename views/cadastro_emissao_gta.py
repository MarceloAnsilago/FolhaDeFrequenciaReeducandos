from datetime import date
from io import BytesIO
from pathlib import Path

import streamlit as st
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from streamlit.errors import StreamlitAPIException
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.units import mm
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas

PERMISSAO_COLUNAS = [
    ("novo", "NOVO"),
    ("editar", "EDITAR"),
    ("cancelar", "CANCELAR"),
    ("consultar", "CONSULTAR"),
]

PERMISSAO_ITENS = [
    (1, "CADASTRO PESSOA FISICA"),
    (2, "CADASTRO PROPRIEDADE RURAL"),
    (3, "CADASTRO FICHA DE BOVIDEOS"),
    (4, "CADASTRO DE LOGRADOURO"),
    (5, "SETOR MUNICIPIO/EPIDEMIOLOGICO"),
    (6, "CAD. FRIGORIFICO/MATADOURO..."),
    (7, "CAD. REVENDEDOR DE VACINAS"),
    (8, "DECLARACAO DE VACINAS"),
    (9, "GTA / TTRB / D. CONS./ ETC e outras\nformas de entrada e saida de bovideos"),
]

PERMISSAO_DEFAULTS = {
    1: {"novo": True, "editar": True, "cancelar": False, "consultar": True},
    2: {"novo": True, "editar": True, "cancelar": False, "consultar": True},
    3: {"novo": True, "editar": True, "cancelar": False, "consultar": True},
    4: {"novo": False, "editar": False, "cancelar": False, "consultar": True},
    5: {"novo": False, "editar": False, "cancelar": False, "consultar": True},
    6: {"novo": False, "editar": False, "cancelar": False, "consultar": True},
    7: {"novo": False, "editar": False, "cancelar": False, "consultar": True},
    8: {"novo": True, "editar": True, "cancelar": False, "consultar": True},
    9: {"novo": True, "editar": True, "cancelar": False, "consultar": True},
}


def _fmt_date(value) -> str:
    if isinstance(value, date):
        return value.strftime("%d/%m/%Y")
    if value:
        return str(value)
    return ""


def _wrap_lines(text: str, font_name: str, font_size: int, max_width: float):
    words = (text or "").split()
    if not words:
        return [""]

    lines = []
    current = words[0]
    for word in words[1:]:
        candidate = f"{current} {word}"
        if stringWidth(candidate, font_name, font_size) <= max_width:
            current = candidate
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines


def _draw_labeled_cell(c: canvas.Canvas, x: float, y: float, w: float, h: float, label: str, value: str):
    c.setLineWidth(0.6)
    c.rect(x, y, w, h)

    c.setFont("Helvetica", 5.2)
    c.drawString(x + 1.3 * mm, y + h - 3.4 * mm, f"{label}:")

    max_w = w - 2.6 * mm
    value_lines = _wrap_lines(value, "Helvetica-Bold", 9, max_w)
    c.setFont("Helvetica-Bold", 9)

    line_y = y + h - 7.2 * mm
    for line in value_lines[:2]:
        c.drawString(x + 1.3 * mm, line_y, line)
        line_y -= 3.8 * mm


def _set_docx_labeled_cell(cell, label: str, value: str, value_size: int = 11):
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p_label = cell.paragraphs[0]
    p_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_label.paragraph_format.space_before = Pt(0)
    p_label.paragraph_format.space_after = Pt(0)
    p_label.paragraph_format.line_spacing = 1
    run_label = p_label.add_run(f"{label}:")
    run_label.font.size = Pt(7)

    p_val = cell.add_paragraph(value or "")
    p_val.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_val.paragraph_format.space_before = Pt(0)
    p_val.paragraph_format.space_after = Pt(0)
    p_val.paragraph_format.line_spacing = 1
    run_val = p_val.runs[0] if p_val.runs else p_val.add_run("")
    run_val.bold = True
    run_val.font.size = Pt(value_size)


def _mm_to_twips(value_mm: float) -> int:
    return int(value_mm * 56.6929)


def _set_table_fixed_width(table, width_mm: float):
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(_mm_to_twips(width_mm)))

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def _set_cell_width(cell, width_mm: float):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(_mm_to_twips(width_mm)))


def _add_docx_header(section, logo_path: Path):
    header = section.header

    p_logo = header.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_before = Pt(0)
    p_logo.paragraph_format.space_after = Pt(2)
    if logo_path.exists():
        run = p_logo.add_run()
        run.add_picture(str(logo_path), width=Mm(24))

    p = header.add_paragraph("GOVERNO DO ESTADO DE RONDONIA")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(8.5)

    p = header.add_paragraph("Agencia de Defesa Sanitaria Agrosilvopastoril do Estado de Rondonia - IDARON")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.runs[0].font.size = Pt(7)


def build_docx_cadastro_gta(data: dict, logo_path: Path) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.top_margin = Mm(36)
    section.bottom_margin = Mm(15)
    section.header_distance = Mm(6)

    _add_docx_header(section, logo_path)

    content_w = 170  # 210 - 20 - 20 mm

    p = doc.add_paragraph("CADASTRO DE SERVIDOR PARA EMISSAO DE GTA")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)

    table = doc.add_table(rows=12, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_fixed_width(table, content_w)

    col_w = [26, 98, 46]  # total = 170mm
    for row in table.rows:
        _set_cell_width(row.cells[0], col_w[0])
        _set_cell_width(row.cells[1], col_w[1])
        _set_cell_width(row.cells[2], col_w[2])

    row_h = [9, 7, 7, 8, 8, 8, 8, 8, 8, 8, 8, 11]
    for idx, row in enumerate(table.rows):
        row.height = Mm(row_h[idx])
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    photo_cell = table.cell(0, 0)
    for row_idx in range(1, 5):
        photo_cell = photo_cell.merge(table.cell(row_idx, 0))
    photo_cell.text = "FOTO\n3x4"
    photo_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if len(photo_cell.paragraphs) > 1:
        photo_cell.paragraphs[1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    photo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    _set_docx_labeled_cell(table.cell(0, 1).merge(table.cell(0, 2)), "NOME", data.get("nome", ""))
    _set_docx_labeled_cell(table.cell(1, 1).merge(table.cell(1, 2)), "CARGO", data.get("cargo", ""))
    _set_docx_labeled_cell(table.cell(2, 1).merge(table.cell(2, 2)), "FORMACAO", data.get("formacao", ""))
    _set_docx_labeled_cell(table.cell(3, 1).merge(table.cell(3, 2)), "MATRICULA", data.get("matricula", ""))
    _set_docx_labeled_cell(table.cell(4, 1), "RG", data.get("rg", ""))
    _set_docx_labeled_cell(table.cell(4, 2), "CPF", data.get("cpf", ""))

    _set_docx_labeled_cell(
        table.cell(5, 0).merge(table.cell(5, 1)),
        "ORGAO DE ORIGEM",
        data.get("orgao_origem", ""),
    )
    _set_docx_labeled_cell(table.cell(5, 2), "DATA DA EMISSAO", data.get("data_emissao", ""))

    _set_docx_labeled_cell(
        table.cell(6, 0).merge(table.cell(6, 1)).merge(table.cell(6, 2)),
        "REGIONAL",
        data.get("regional", ""),
    )

    _set_docx_labeled_cell(
        table.cell(7, 0).merge(table.cell(7, 1)),
        "UNIDADE DE LOTACAO",
        data.get("unidade_lotacao", ""),
    )
    _set_docx_labeled_cell(table.cell(7, 2), "DATA DE LOTACAO", data.get("data_lotacao", ""))

    _set_docx_labeled_cell(
        table.cell(8, 0).merge(table.cell(8, 1)).merge(table.cell(8, 2)),
        "AUTORIZADO PARA TRANSITO",
        data.get("autorizado_transito", ""),
    )
    _set_docx_labeled_cell(
        table.cell(9, 0).merge(table.cell(9, 1)).merge(table.cell(9, 2)),
        "MUNIC/EST. AUTORIZADO",
        data.get("municipio_estado", ""),
    )
    _set_docx_labeled_cell(
        table.cell(10, 0).merge(table.cell(10, 1)).merge(table.cell(10, 2)),
        "ESPECIES AUTORIZADAS",
        data.get("especies", ""),
    )
    _set_docx_labeled_cell(
        table.cell(11, 0).merge(table.cell(11, 1)).merge(table.cell(11, 2)),
        "OUTROS DOCUMENTOS",
        data.get("outros_documentos", ""),
    )

    doc.add_paragraph("")
    p = doc.add_paragraph("DEMONSTRATIVO DE ASSINATURAS DO SERVIDOR")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    doc.add_paragraph("")
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_table.autofit = False
    sig_table.style = "Table Grid"
    _set_table_fixed_width(sig_table, content_w)

    for row in sig_table.rows:
        _set_cell_width(row.cells[0], content_w / 2)
        _set_cell_width(row.cells[1], content_w / 2)
    for row in sig_table.rows:
        row.height = Mm(30)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    # 1) Assinatura (superior esquerda)
    cell = sig_table.cell(0, 0)
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("______________________________")
    p = cell.add_paragraph("(ASSINATURA DO SERVIDOR)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 2) Assinatura (superior direita)
    cell = sig_table.cell(0, 1)
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("______________________________")
    p = cell.add_paragraph("(ASSINATURA DO SERVIDOR)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 3) Assinatura (inferior esquerda)
    cell = sig_table.cell(1, 0)
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("______________________________")
    p = cell.add_paragraph("(ASSINATURA DO SERVIDOR)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 4) Ciente + carimbo/assinatura (inferior direita)
    cell = sig_table.cell(1, 1)
    cell.text = ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("CIENTE:")
    run.bold = True
    run.font.size = Pt(14)
    p = cell.add_paragraph("______________________________")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = cell.add_paragraph("(CARIMBO E ASSINATURA)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def build_pdf_cadastro_gta(data: dict, logo_path: Path) -> bytes:
    buffer = BytesIO()
    page_width, page_height = A4
    c = canvas.Canvas(buffer, pagesize=A4)

    # Moldura externa
    page_margin = 8 * mm
    c.setLineWidth(0.7)
    c.rect(page_margin, page_margin, page_width - 2 * page_margin, page_height - 2 * page_margin)

    x = 20 * mm
    w = page_width - 40 * mm

    # Cabecalho superior
    y_top = page_height - 18 * mm
    if logo_path.exists():
        logo = ImageReader(str(logo_path))
        img_w, img_h = logo.getSize()
        max_size = 14 * mm
        scale = min(max_size / img_w, max_size / img_h)
        draw_w = img_w * scale
        draw_h = img_h * scale
        c.drawImage(
            logo,
            x + (w - draw_w) / 2,
            y_top - draw_h,
            width=draw_w,
            height=draw_h,
            mask="auto",
        )

    c.setFont("Helvetica-Bold", 9)
    c.drawCentredString(x + w / 2, y_top - 16 * mm, "GOVERNO DO ESTADO DE RONDONIA")
    c.setFont("Helvetica", 6.4)
    c.drawCentredString(
        x + w / 2,
        y_top - 19.8 * mm,
        "Agencia de Defesa Sanitaria Agrosilvopastoril do Estado de Rondonia - IDARON",
    )

    title_h = 9 * mm
    title_y = y_top - 30 * mm
    c.rect(x, title_y, w, title_h)
    c.setFont("Helvetica-Bold", 14)
    c.drawCentredString(x + w / 2, title_y + 2.9 * mm, "CADASTRO DE SERVIDOR PARA EMISSAO DE GTA")

    table_y = title_y - 5 * mm

    row_specs = [
        ("NOME", "nome", 9 * mm, None),
        ("CARGO", "cargo", 7 * mm, None),
        ("FORMACAO", "formacao", 7 * mm, None),
        ("MATRICULA", "matricula", 8 * mm, None),
        (("RG", "rg"), ("CPF", "cpf"), 8 * mm, 0.6),
        (("ORGAO DE ORIGEM", "orgao_origem"), ("DATA DA EMISSAO", "data_emissao"), 8 * mm, 0.76),
        ("REGIONAL", "regional", 8 * mm, None),
        (("UNIDADE DE LOTACAO", "unidade_lotacao"), ("DATA DE LOTACAO", "data_lotacao"), 8 * mm, 0.76),
        ("AUTORIZADO PARA TRANSITO", "autorizado_transito", 8 * mm, None),
        ("MUNIC/EST. AUTORIZADO", "municipio_estado", 8 * mm, None),
        ("ESPECIES AUTORIZADAS", "especies", 8 * mm, None),
        ("OUTROS DOCUMENTOS", "outros_documentos", 12 * mm, None),
    ]

    # Primeiras linhas com recuo para reservar o quadro da foto.
    photo_specs = row_specs[:5]
    remaining_specs = row_specs[5:]
    photo_w = 24 * mm
    photo_h = sum(spec[2] for spec in photo_specs)

    y = table_y
    photo_bottom_y = y - photo_h
    c.setLineWidth(0.6)
    c.rect(x, photo_bottom_y, photo_w, photo_h)
    c.setFont("Helvetica", 7)
    c.drawCentredString(x + (photo_w / 2), photo_bottom_y + (photo_h / 2), "FOTO")
    c.setFont("Helvetica", 6)
    c.drawCentredString(x + (photo_w / 2), photo_bottom_y + (photo_h / 2) - 4.2 * mm, "3x4")

    fields_x = x + photo_w
    fields_w = w - photo_w
    for spec in photo_specs:
        height = spec[2]
        y -= height

        split = spec[3]
        if split is None:
            label, key = spec[0], spec[1]
            _draw_labeled_cell(c, fields_x, y, fields_w, height, label, data.get(key, ""))
        else:
            left, right = spec[0], spec[1]
            left_w = fields_w * split
            right_w = fields_w - left_w
            _draw_labeled_cell(c, fields_x, y, left_w, height, left[0], data.get(left[1], ""))
            _draw_labeled_cell(c, fields_x + left_w, y, right_w, height, right[0], data.get(right[1], ""))

    for spec in remaining_specs:
        height = spec[2]
        y -= height

        split = spec[3]
        if split is None:
            label, key = spec[0], spec[1]
            _draw_labeled_cell(c, x, y, w, height, label, data.get(key, ""))
        else:
            left, right = spec[0], spec[1]
            left_w = w * split
            right_w = w - left_w
            _draw_labeled_cell(c, x, y, left_w, height, left[0], data.get(left[1], ""))
            _draw_labeled_cell(c, x + left_w, y, right_w, height, right[0], data.get(right[1], ""))

    y -= 12 * mm
    c.rect(x, y, w, 8 * mm)
    c.setFont("Helvetica-Bold", 11)
    c.drawCentredString(x + w / 2, y + 2.5 * mm, "DEMONSTRATIVO DE ASSINATURAS DO SERVIDOR")

    # Assinaturas inferiores
    sig_w = 70 * mm
    sig1_x = x
    sig1_y = y - 34 * mm
    sig2_x = x + w - sig_w
    sig2_y = y - 58 * mm
    sig3_x = x
    sig3_y = y - 82 * mm
    sig4_x = x + w - sig_w
    sig4_y = y - 106 * mm

    for sx, sy in [(sig1_x, sig1_y), (sig2_x, sig2_y), (sig3_x, sig3_y)]:
        c.line(sx, sy, sx + sig_w, sy)
        c.setFont("Helvetica", 5.4)
        c.drawCentredString(sx + sig_w / 2, sy - 3.3 * mm, "(ASSINATURA DO SERVIDOR)")

    c.setFont("Helvetica-Bold", 16)
    c.drawString(sig4_x + 2 * mm, sig4_y + 15 * mm, "CIENTE:")
    c.line(sig4_x, sig4_y, sig4_x + sig_w, sig4_y)
    c.setFont("Helvetica", 5.4)
    c.drawCentredString(sig4_x + sig_w / 2, sig4_y - 3.3 * mm, "(CARIMBO E ASSINATURA)")

    c.showPage()
    c.save()
    buffer.seek(0)
    return buffer.read()


def build_pdf_permissoes_gta(data: dict, logo_path: Path, permissoes: dict) -> bytes:
    buffer = BytesIO()
    page_width, page_height = landscape(A4)
    c = canvas.Canvas(buffer, pagesize=(page_width, page_height))

    page_margin = 8 * mm
    c.setLineWidth(0.7)
    c.rect(page_margin, page_margin, page_width - 2 * page_margin, page_height - 2 * page_margin)

    x = 18 * mm
    w = page_width - 36 * mm
    y_top = page_height - 12 * mm

    if logo_path.exists():
        logo = ImageReader(str(logo_path))
        img_w, img_h = logo.getSize()
        max_size = 11 * mm
        scale = min(max_size / img_w, max_size / img_h)
        draw_w = img_w * scale
        draw_h = img_h * scale
        c.drawImage(
            logo,
            x + (w - draw_w) / 2,
            y_top - draw_h,
            width=draw_w,
            height=draw_h,
            mask="auto",
        )

    c.setFont("Helvetica-Bold", 8.5)
    c.drawCentredString(x + w / 2, y_top - 13 * mm, "GOVERNO DO ESTADO DE RONDONIA")
    c.setFont("Helvetica", 6.2)
    c.drawCentredString(
        x + w / 2,
        y_top - 16.5 * mm,
        "Agencia de Defesa Sanitaria Agrosilvopastoril do Estado de Rondonia - IDARON",
    )
    c.setFont("Helvetica-Bold", 11)
    c.drawCentredString(x + w / 2, y_top - 24.5 * mm, "NIVEIS DE PERMISSOES PARA USUARIOS DO SISTEMA SISIDARON")

    info_y = y_top - 33 * mm
    unidade_lotacao = (data.get("unidade_lotacao", "") or "").strip()
    c.setFont("Helvetica-Bold", 9)
    c.drawString(x + 2 * mm, info_y, f"Nome do Servidor:  {data.get('nome', '')}")
    c.drawString(x + 112 * mm, info_y, f"Unidade de Lotacao: {unidade_lotacao}")

    c.drawString(x + 2 * mm, info_y - 8 * mm, f"Funcao:  {data.get('cargo', '')}")
    c.drawString(x + 94 * mm, info_y - 8 * mm, f"CPF: {data.get('cpf', '')}")
    c.drawString(x + 150 * mm, info_y - 8 * mm, f"Matricula: {data.get('matricula', '')}")

    table_w = 200 * mm
    table_x = x + (w - table_w) / 2
    table_top = info_y - 14 * mm
    col_ws = [18 * mm, 74 * mm, 27 * mm, 27 * mm, 27 * mm, 27 * mm]
    body_row_hs = []
    for _, item_nome in PERMISSAO_ITENS:
        body_row_hs.append(5.4 * mm)  # linha com dados
        body_row_hs.append(5.4 * mm if "\n" in item_nome else 3.1 * mm)  # linha em branco abaixo
    row_hs = [7.2 * mm, 6.2 * mm] + body_row_hs

    c.setFillColorRGB(0.05, 0.1, 0.85)
    c.rect(table_x, table_top - row_hs[0], table_w, row_hs[0], stroke=1, fill=1)
    c.rect(table_x, table_top - row_hs[0] - row_hs[1], table_w, row_hs[1], stroke=1, fill=1)
    c.setFillColorRGB(0, 0, 0)

    total_h = sum(row_hs)
    c.setLineWidth(0.8)
    c.rect(table_x, table_top - total_h, table_w, total_h)

    y_line = table_top
    for h in row_hs:
        y_line -= h
        c.line(table_x, y_line, table_x + table_w, y_line)

    x_line = table_x
    for w_col in col_ws:
        x_line += w_col
        c.line(x_line, table_top - total_h, x_line, table_top)

    x_perm_start = table_x + col_ws[0] + col_ws[1]
    c.setLineWidth(1.0)
    c.line(x_perm_start, table_top - row_hs[0], x_perm_start + sum(col_ws[2:]), table_top - row_hs[0])
    c.setLineWidth(0.6)

    c.setFillColorRGB(1, 1, 1)
    c.setFont("Helvetica-Bold", 10)
    c.drawCentredString(table_x + col_ws[0] / 2, table_top - 5.1 * mm, "ITEM")
    c.drawCentredString(table_x + col_ws[0] + col_ws[1] / 2, table_top - 5.1 * mm, "DOCUMENTO")
    c.drawCentredString(x_perm_start + sum(col_ws[2:]) / 2, table_top - 3.5 * mm, "PERMISSOES")

    y_head2 = table_top - row_hs[0] - 4.6 * mm
    for idx, (_, label) in enumerate(PERMISSAO_COLUNAS):
        x_c = x_perm_start + sum(col_ws[2 : 2 + idx]) + col_ws[2 + idx] / 2
        c.drawCentredString(x_c, y_head2, label)
    c.setFillColorRGB(0, 0, 0)

    def _draw_cell_lines(text: str, x_left: float, x_right: float, y_top_row: float, row_h: float, align: str = "left"):
        lines = text.split("\n")
        font_size = 9
        line_step = 11  # points
        y_center = y_top_row - (row_h / 2)
        y_first = y_center + ((len(lines) - 1) * line_step / 2) - (font_size * 0.32)
        for idx_line, line in enumerate(lines):
            y_line = y_first - (idx_line * line_step)
            if align == "center":
                c.drawCentredString((x_left + x_right) / 2, y_line, line)
            else:
                c.drawString(x_left + 1.2 * mm, y_line, line)

    y_cursor = table_top - row_hs[0] - row_hs[1]
    c.setFont("Helvetica", 9)
    body_row_idx = 2
    for item_num, item_nome in PERMISSAO_ITENS:
        data_h = row_hs[body_row_idx]
        blank_h = row_hs[body_row_idx + 1]
        body_row_idx += 2

        item_lines = item_nome.split("\n")
        first_line = item_lines[0]
        second_line = item_lines[1] if len(item_lines) > 1 else ""

        _draw_cell_lines(
            str(item_num),
            table_x,
            table_x + col_ws[0],
            y_cursor,
            data_h,
            align="center",
        )
        _draw_cell_lines(
            first_line,
            table_x + col_ws[0],
            table_x + col_ws[0] + col_ws[1],
            y_cursor,
            data_h,
            align="left",
        )

        checks = permissoes.get(item_num, {})
        for col_idx, (perm_key, _) in enumerate(PERMISSAO_COLUNAS):
            if checks.get(perm_key):
                x_left = x_perm_start + sum(col_ws[2 : 2 + col_idx])
                x_right = x_left + col_ws[2 + col_idx]
                _draw_cell_lines("X", x_left, x_right, y_cursor, data_h, align="center")

        if second_line:
            _draw_cell_lines(
                second_line,
                table_x + col_ws[0],
                table_x + col_ws[0] + col_ws[1],
                y_cursor - data_h,
                blank_h,
                align="left",
            )

        y_cursor -= data_h + blank_h

    c.setFont("Helvetica", 10)
    c.drawString(table_x + 1 * mm, table_top - total_h - 8 * mm, "Marcar com X a permissao que o servidor tera acesso no SISIDARON")

    sig_y = table_top - total_h - 43 * mm
    sig_w = 70 * mm
    left_sig_x = table_x + 20 * mm
    right_sig_x = table_x + table_w - 20 * mm - sig_w

    c.line(left_sig_x, sig_y, left_sig_x + sig_w, sig_y)
    c.drawCentredString(left_sig_x + sig_w / 2, sig_y - 4.2 * mm, "Assinatura do Funcionario")
    c.line(right_sig_x, sig_y, right_sig_x + sig_w, sig_y)
    c.drawCentredString(right_sig_x + sig_w / 2, sig_y - 4.2 * mm, "Assinatura do Chefe de ULSAV")

    via_y = sig_y - 13 * mm
    c.setFont("Helvetica-Bold", 12)
    c.drawString(table_x + 1 * mm, via_y, "1a via: SEINF/GID SA")
    c.drawCentredString(table_x + table_w / 2, via_y, "2a via: REGIONAL")
    c.drawRightString(table_x + table_w - 1 * mm, via_y, "3a via: ULSAV")

    c.showPage()
    c.save()
    buffer.seek(0)
    return buffer.read()


def render_cadastro_emissao_gta():
    st.session_state.setdefault("gta_autorizado_transito", "Intramunicipal, Intermunicipal, Intraestadual")
    st.session_state.setdefault("gta_municipio_estado", "Todo o Estado")
    st.session_state.setdefault("gta_especies", "Bovinos/Ovinos/Caprinos/Suinos/Equideos")
    st.session_state.setdefault(
        "gta_outros_documentos",
        "Todos os documentos, exceto CIS-E, CTC, GTR, Auto de Infracao",
    )
    for item_num, _ in PERMISSAO_ITENS:
        defaults_item = PERMISSAO_DEFAULTS.get(item_num, {})
        for perm_key, _ in PERMISSAO_COLUNAS:
            st.session_state.setdefault(f"gta_perm_{item_num}_{perm_key}", defaults_item.get(perm_key, False))

    col_left, col_mid, col_right = st.columns([1, 2, 1])
    with col_mid:
        st.title("Cadastro de Emissao de GTA")

        with st.form("form_cadastro_emissao_gta"):
            st.text_input("Nome", key="gta_nome")
            st.text_input("Cargo", key="gta_cargo")
            st.text_input("Formacao", key="gta_formacao")
            st.text_input("Matricula", key="gta_matricula")

            col_doc = st.columns(2)
            with col_doc[0]:
                st.text_input("RG", key="gta_rg")
            with col_doc[1]:
                st.text_input("CPF", key="gta_cpf")

            col_orgao = st.columns(2)
            with col_orgao[0]:
                st.text_input("Orgao de origem", key="gta_orgao_origem")
            with col_orgao[1]:
                st.date_input("Data da emissao", key="gta_data_emissao", value=date.today())

            st.text_input("Regional", key="gta_regional")

            col_lot = st.columns(2)
            with col_lot[0]:
                st.text_input("Unidade de lotacao", key="gta_unidade_lotacao")
            with col_lot[1]:
                st.date_input("Data de lotacao", key="gta_data_lotacao", value=date.today())

            st.text_input("Autorizado para transito", key="gta_autorizado_transito")
            st.text_input("Munic./Est. autorizado", key="gta_municipio_estado")
            st.text_input("Especies autorizadas", key="gta_especies")
            st.text_area("Outros documentos", key="gta_outros_documentos", height=70)

            submit = st.form_submit_button("Gerar Arquivos")

        if submit:
            logo_path = Path(__file__).resolve().parents[1] / "assets" / "logo_ro_horizontal.JPG"
            data = {
                "nome": st.session_state.get("gta_nome", ""),
                "cargo": st.session_state.get("gta_cargo", ""),
                "formacao": st.session_state.get("gta_formacao", ""),
                "matricula": st.session_state.get("gta_matricula", ""),
                "rg": st.session_state.get("gta_rg", ""),
                "cpf": st.session_state.get("gta_cpf", ""),
                "orgao_origem": st.session_state.get("gta_orgao_origem", ""),
                "data_emissao": _fmt_date(st.session_state.get("gta_data_emissao")),
                "regional": st.session_state.get("gta_regional", ""),
                "unidade_lotacao": st.session_state.get("gta_unidade_lotacao", ""),
                "data_lotacao": _fmt_date(st.session_state.get("gta_data_lotacao")),
                "autorizado_transito": st.session_state.get("gta_autorizado_transito", ""),
                "municipio_estado": st.session_state.get("gta_municipio_estado", ""),
                "especies": st.session_state.get("gta_especies", ""),
                "outros_documentos": st.session_state.get("gta_outros_documentos", ""),
            }

            pdf_bytes = build_pdf_cadastro_gta(data, logo_path)
            docx_bytes = build_docx_cadastro_gta(data, logo_path)
            st.session_state["cadastro_gta_pdf"] = pdf_bytes
            st.session_state["cadastro_gta_docx"] = docx_bytes

            output_dir = Path(__file__).resolve().parents[1] / "pdf"
            output_dir.mkdir(parents=True, exist_ok=True)
            output_path = output_dir / "cadastro_emissao_gta.pdf"
            output_docx_path = output_dir / "cadastro_emissao_gta.docx"
            output_path.write_bytes(pdf_bytes)
            output_docx_path.write_bytes(docx_bytes)
            st.session_state["cadastro_gta_pdf_path"] = output_path
            st.session_state["cadastro_gta_docx_path"] = output_docx_path
            st.success("Arquivos PDF e DOCX gerados e salvos.")

        if "cadastro_gta_pdf" in st.session_state:
            st.markdown("### Pagina de impressao")
            pdf_bytes = st.session_state["cadastro_gta_pdf"]
            try:
                st.pdf(pdf_bytes)
            except StreamlitAPIException:
                st.info(
                    "Pre-visualizacao de PDF indisponivel neste ambiente. "
                    "Para habilitar, instale: pip install streamlit[pdf]"
                )

            st.download_button(
                "Baixar PDF",
                data=pdf_bytes,
                file_name="cadastro_emissao_gta.pdf",
                mime="application/pdf",
            )
            if "cadastro_gta_docx" in st.session_state:
                st.download_button(
                    "Baixar DOCX (OnlyOffice/LibreOffice)",
                    data=st.session_state["cadastro_gta_docx"],
                    file_name="cadastro_emissao_gta.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            if "cadastro_gta_pdf_path" in st.session_state:
                st.caption(f"Salvo em: {st.session_state['cadastro_gta_pdf_path']}")
            if "cadastro_gta_docx_path" in st.session_state:
                st.caption(f"DOCX salvo em: {st.session_state['cadastro_gta_docx_path']}")

        st.divider()
        st.subheader("Formulario de Permissoes SISIDARON (PDF Paisagem)")
        st.caption("Este PDF usa os mesmos dados do cadastro acima (nome, funcao, CPF, matricula e lotacao).")

        with st.form("form_permissoes_sisidaron"):
            for item_num, item_nome in PERMISSAO_ITENS:
                st.markdown(f"**{item_num}. {item_nome.replace(chr(10), ' / ')}**")
                cols = st.columns(4)
                for col_idx, (perm_key, perm_label) in enumerate(PERMISSAO_COLUNAS):
                    with cols[col_idx]:
                        st.checkbox(perm_label, key=f"gta_perm_{item_num}_{perm_key}")

            submit_permissoes = st.form_submit_button("Gerar PDF de Permissoes")

        if submit_permissoes:
            logo_path = Path(__file__).resolve().parents[1] / "assets" / "logo_ro_horizontal.JPG"
            data = {
                "nome": st.session_state.get("gta_nome", ""),
                "cargo": st.session_state.get("gta_cargo", ""),
                "matricula": st.session_state.get("gta_matricula", ""),
                "cpf": st.session_state.get("gta_cpf", ""),
                "unidade_lotacao": st.session_state.get("gta_unidade_lotacao", ""),
            }
            permissoes = {}
            for item_num, _ in PERMISSAO_ITENS:
                permissoes[item_num] = {}
                for perm_key, _ in PERMISSAO_COLUNAS:
                    permissoes[item_num][perm_key] = st.session_state.get(f"gta_perm_{item_num}_{perm_key}", False)

            pdf_permissoes = build_pdf_permissoes_gta(data, logo_path, permissoes)
            st.session_state["cadastro_gta_permissoes_pdf"] = pdf_permissoes

            output_dir = Path(__file__).resolve().parents[1] / "pdf"
            output_dir.mkdir(parents=True, exist_ok=True)
            output_path = output_dir / "cadastro_emissao_gta_permissoes.pdf"
            output_path.write_bytes(pdf_permissoes)
            st.session_state["cadastro_gta_permissoes_pdf_path"] = output_path
            st.success("PDF de permissoes gerado e salvo.")

        if "cadastro_gta_permissoes_pdf" in st.session_state:
            st.markdown("### Pagina de impressao - Permissoes")
            pdf_bytes = st.session_state["cadastro_gta_permissoes_pdf"]
            try:
                st.pdf(pdf_bytes)
            except StreamlitAPIException:
                st.info(
                    "Pre-visualizacao de PDF indisponivel neste ambiente. "
                    "Para habilitar, instale: pip install streamlit[pdf]"
                )

            st.download_button(
                "Baixar PDF de Permissoes",
                data=pdf_bytes,
                file_name="cadastro_emissao_gta_permissoes.pdf",
                mime="application/pdf",
            )
            if "cadastro_gta_permissoes_pdf_path" in st.session_state:
                st.caption(f"Salvo em: {st.session_state['cadastro_gta_permissoes_pdf_path']}")
