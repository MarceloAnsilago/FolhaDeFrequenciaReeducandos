import base64
import locale
import re
from datetime import datetime
from io import BytesIO
from pathlib import Path

import pandas as pd
import streamlit as st
import streamlit.components.v1 as components
from streamlit_option_menu import option_menu
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.units import mm
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None  # type: ignore

try:
    from docx import Document
except ImportError:
    Document = None  # type: ignore

from pdf.cabecalho import desenhar_cabecalho
from pdf.corpo import desenhar_tabela
from pdf.relatorio import gerar_relatorio_cabecalho
from pdf.rodape import desenhar_rodape
from pdf.sugesp import gerar_pdf_sugesp


st.set_page_config(page_title="Folha de Ponto de Reeducandos", page_icon="📄", layout="wide")

MESES = {
    "JANEIRO": 1,
    "FEVEREIRO": 2,
    "MARÇO": 3,
    "ABRIL": 4,
    "MAIO": 5,
    "JUNHO": 6,
    "JULHO": 7,
    "AGOSTO": 8,
    "SETEMBRO": 9,
    "OUTUBRO": 10,
    "NOVEMBRO": 11,
    "DEZEMBRO": 12,
}

ANO_INICIO = 2025
ANOS_OPCOES = list(range(ANO_INICIO, ANO_INICIO + 11))
ANO_ATUAL = datetime.now().year

# session state defaults
DEFAULTS = {
    "secretaria": "",
    "reeducando": "",
    "funcao": "",
    "data_inclusao": "",
    "municipio": "",
    "cpf": "",
    "banco": "",
    "agencia": "",
    "conta": "",
    "tipo_conta": "",
    "endereco": "",
    "cep": "",
    "telefone": "",
    "data_preenchimento": "__/__/____",
    "rodape_titulo": "ULSAV - UNIDADE LOCAL DE SANIDADE ANIMAL E VEGETAL",
    "rodape_endereco": "Av. São Paulo, 436 – Bairro Centro",
    "rodape_fone": "Fone/Fax: (69) 3642-1026/8479-9229",
    "rodape_cep": "CEP 76.932-000 – São Miguel do Guaporé/RO",
    "rodape_email": "saomiguel@idaron.ro.gov.br",
    "mes_label": list(MESES.keys())[0],
    "ano": ANO_ATUAL if ANO_ATUAL in ANOS_OPCOES else ANOS_OPCOES[0],
    "he": "07:30",
    "hs": "13:30",
    "feriados_texto": "",
    "sugesp_unidade": "SUPERINTENDENCIA DE GESTAO DOS GASTOS PUBLICOS ADMINISTRATIVOS - SUGESP",
    "sugesp_sub_unidade": "",
    "sugesp_setor_lotacao": "",
    "sugesp_servidor": "",
    "sugesp_matricula": "",
    "sugesp_sigla": "",
    "sugesp_cargo": "",
    "sugesp_endereco": "",
    "sugesp_cep": "",
    "sugesp_telefone": "",
    "sugesp_email": "",
    "sugesp_cpf": "",
    "sugesp_data_preenchimento": "__/__/____",
    "sugesp_mes_label": list(MESES.keys())[0],
    "sugesp_ano": ANO_ATUAL if ANO_ATUAL in ANOS_OPCOES else ANOS_OPCOES[0],
    "sugesp_he": "07:30",
    "sugesp_hs": "13:30",
    "sugesp_feriados_texto": "",
}

for chave, valor in DEFAULTS.items():
    st.session_state.setdefault(chave, valor)
# flags de controle do upload (para não sobrescrever após a primeira aplicação)
st.session_state.setdefault("_upload_aplicado", False)
st.session_state.setdefault("_ultimo_upload", "")
st.session_state.setdefault("_sugesp_upload_aplicado", False)
st.session_state.setdefault("_sugesp_ultimo_upload", "")


def _safe_index(options, value, default=0):
    try:
        return options.index(value)
    except Exception:
        return default


def parse_feriados_text(texto):
    feriados_dict = {}
    erros = []
    for raw_bloco in texto.split(","):
        bloco = raw_bloco.strip()
        if not bloco:
            continue
        if "-" not in bloco:
            erros.append(f'"{bloco}" (faltou "-")')
            continue
        dia_str, _, nome = bloco.partition("-")
        dia_str = dia_str.strip()
        nome = nome.strip()
        try:
            dia = int(dia_str)
        except ValueError:
            erros.append(f'"{bloco}" (dia inválido)')
            continue
        if not (1 <= dia <= 31):
            erros.append(f'"{bloco}" (dia fora de 1-31)')
            continue
        if not nome:
            erros.append(f'"{bloco}" (descrição vazia)')
            continue
        feriados_dict[dia] = nome
    return feriados_dict, erros


def _clean_text(texto: str) -> str:
    # normaliza espaços e sobe para maiúsculas para regex
    return re.sub(r"\s+", " ", texto).strip()


def _parse_campos(texto: str) -> dict:
    """Extrai campos do texto plano (PDF ou DOCX)."""
    norm = _clean_text(texto).upper()
    campos = {}

    def normaliza_data(txt: str) -> str:
        if "_" in txt:
            return "__/__/____"
        bruto = txt.strip()
        digitos = re.sub(r"\D", "", bruto)
        if len(digitos) == 8:
            return f"{digitos[:2]}/{digitos[2:4]}/{digitos[4:]}"
        if not digitos:
            return "__/__/____"
        return "__/__/____"

    def normaliza_tipo_conta(txt: str) -> str:
        t = txt.upper()
        t_norm = re.sub(r"\s+", " ", t)
        # prioridade: opção marcada com (X)
        if re.search(r"\(\s*X\s*\)\s*CORRENTE", t_norm):
            return "Corrente"
        if re.search(r"\(\s*X\s*\)\s*SAL[AA]RIO", t_norm):
            return "Salário"
        if re.search(r"\(\s*X\s*\)\s*POUPAN[CC]A", t_norm):
            return "Poupança"
        # fallback por palavra-chave
        if "POUP" in t:
            return "Poupança"
        if "SAL" in t:  # SALÁRIO
            return "Salário"
        if "CORRENTE" in t:
            return "Corrente"
        return txt

    def pega(padrao, grupo=1):
        m = re.search(padrao, norm)
        return m.group(grupo).strip() if m else ""

    campos["secretaria"] = pega(r"SECRETARIA:\s*(.+?)(?:\s+ANO:|$)")
    campos["ano"] = pega(r"ANO:\s*(\d{4})")
    campos["reeducando"] = pega(r"REEDUCANDO:\s*(.+?)(?:\s+M[ÊE]S:|$)")
    campos["mes_label"] = pega(r"M[ÊE]S:\s*([A-ZÇÃÕ]+)")
    campos["funcao"] = pega(r"FUNÇÃO:\s*(.+?)(?:\s+DATA DA INCLUSÃO:|$)")
    campos["data_inclusao"] = pega(r"DATA DA INCLUS[ÃA]O:\s*([\d/]+)")
    campos["municipio"] = pega(r"MUNIC[IÍ]PIO:\s*(.+?)(?:\s+CPF:|$)")
    campos["cpf"] = pega(r"CPF:\s*([\d.\-]+)")
    campos["banco"] = pega(r"BCO:\s*([A-Z0-9]+)")
    campos["agencia"] = pega(r"AG:\s*([A-Z0-9.\-]+)")
    campos["conta"] = pega(r"CONTA:\s*([A-Z0-9.\-]+)")
    campos["tipo_conta"] = normaliza_tipo_conta(pega(r"TIPO DE CONTA:\s*(.+)"))
    campos["endereco"] = pega(r"ENDEREÇO:\s*(.+?)(?:\s+CEP:|$)")
    campos["cep"] = pega(r"CEP:\s*([\d.\-]+)")
    campos["telefone"] = pega(r"TELEFONE:\s*([0-9\s\-]+)")
    campos["data_preenchimento"] = normaliza_data(pega(r"DATA:\s*([0-9_/]+)"))

    # remove repetições quando o texto veio duplicado no DOCX/PDF
    repetidos = {
        "secretaria": "SECRETARIA:",
        "reeducando": "REEDUCANDO:",
        "funcao": "FUNÇÃO:",
        "municipio": "MUNICÍPIO:",
        "endereco": "ENDEREÇO:",
    }
    for chave, marcador in repetidos.items():
        val = campos.get(chave, "")
        if val:
            up = val.upper()
            idx = up.find(marcador, len(marcador))  # procura segunda ocorrência
            if idx > 0:
                campos[chave] = val[:idx].strip()

    # ano para inteiro, se possível
    try:
        campos["ano"] = int(campos["ano"])
    except Exception:
        campos["ano"] = ""

    # saneia mês para opção conhecida
    mes_upper = campos.get("mes_label", "")
    if mes_upper in MESES:
        campos["mes_label"] = mes_upper
    elif mes_upper:
        # tenta capitalizar conforme dicionário
        for nome in MESES.keys():
            if nome.startswith(mes_upper[:3]):
                campos["mes_label"] = nome
                break

    return campos


def _parse_campos_sugesp(texto: str) -> dict:
    norm = _clean_text(texto).upper()
    campos = {}

    def normaliza_data(txt: str) -> str:
        if "_" in txt:
            return "__/__/____"
        bruto = txt.strip()
        digitos = re.sub(r"\D", "", bruto)
        if len(digitos) == 8:
            return f"{digitos[:2]}/{digitos[2:4]}/{digitos[4:]}"
        if not digitos:
            return "__/__/____"
        return "__/__/____"

    def pega(padrao, grupo=1):
        m = re.search(padrao, norm)
        return m.group(grupo).strip() if m else ""

    def remove_prefixo(valor: str, prefixo_regex: str) -> str:
        if not valor:
            return valor
        return re.sub(rf"^{prefixo_regex}\s*", "", valor, flags=re.IGNORECASE).strip()

    def corta_no_rotulo(valor: str, rotulo_regex: str) -> str:
        if not valor:
            return valor
        m = re.search(rotulo_regex, valor, flags=re.IGNORECASE)
        if not m:
            return valor
        # se o rótulo aparece no meio, mantém apenas a parte antes dele
        if m.start() > 0:
            return valor[: m.start()].strip()
        # se por algum motivo começar com rótulo, remove e retorna o restante
        return re.sub(rf"^{rotulo_regex}\s*", "", valor, flags=re.IGNORECASE).strip()

    campos["sugesp_unidade"] = pega(r"UNIDADE:\s*(.+?)(?:\s+ANO:|$)")
    campos["sugesp_ano"] = pega(r"ANO:\s*(\d{4})")
    campos["sugesp_sub_unidade"] = pega(r"SUB\s*UNIDADE:\s*(.+?)(?:\s+M[ÃŠE]S:|$)")
    campos["sugesp_mes_label"] = pega(r"M[ÃŠE]S:\s*([A-ZÃ‡ÃƒÃ•]+)")
    campos["sugesp_setor_lotacao"] = pega(r"SETOR DE LOTA[ÇC][ÃA]O:\s*(.+?)(?:\s+SERVIDOR:|$)")
    campos["sugesp_servidor"] = pega(r"SERVIDOR:\s*(.+?)(?:\s+MATR[IÍ]CULA:|$)")
    campos["sugesp_matricula"] = pega(r"MATR[IÍ]CULA:\s*([0-9]+)")
    campos["sugesp_sigla"] = pega(r"MATR[IÍ]CULA:\s*[0-9]+\s*([A-Z]{2,4})\b")
    campos["sugesp_cargo"] = pega(r"CARGO:\s*(.+?)(?:\s+DIA|$)")
    campos["sugesp_endereco"] = pega(r"ENDERE[ÇC]O:\s*(.+?)(?:\s+CEP:|$)")
    campos["sugesp_cep"] = pega(r"CEP:\s*([\d.\-]+)")
    campos["sugesp_telefone"] = pega(r"TELEFONE:\s*([0-9()\s\-]+)")
    campos["sugesp_email"] = pega(r"EMAIL:\s*([A-Z0-9._%+\-]+@[A-Z0-9.\-]+\.[A-Z]{2,})")
    campos["sugesp_cpf"] = pega(r"CPF:\s*([\d.\-]+)")
    campos["sugesp_data_preenchimento"] = normaliza_data(pega(r"DATA:\s*([0-9_/]+)"))

    # fallback: pega a sigla da celula abaixo do MES (entre MES e CARGO/DIA)
    sigla = campos.get("sugesp_sigla", "")
    if not sigla or sigla in {"MATR", "MES", "ANO"}:
        m = re.search(r"M[ÊE]S:\s*[A-ZÇÃÕ]+(.*?)(?:CARGO:|DIA)", norm)
        if m:
            trecho = m.group(1)
            tokens = re.findall(r"\b[A-Z]{2,4}\b", trecho)
            for tok in tokens:
                if tok not in {"MATR", "MES", "ANO"}:
                    sigla = tok
            if sigla:
                campos["sugesp_sigla"] = sigla

    repetidos_regex = {
        "sugesp_unidade": r"UNIDADE:",
        "sugesp_sub_unidade": r"SUB\s+UNIDADE:",
        "sugesp_setor_lotacao": r"SETOR DE LOTA[ÇC][ÃA]O:",
        "sugesp_servidor": r"SERVIDOR:",
        "sugesp_cargo": r"CARGO:",
        "sugesp_endereco": r"ENDERE[ÇC]O:",
        "sugesp_telefone": r"TELEFONE:",
        "sugesp_email": r"EMAIL:",
        "sugesp_cpf": r"CPF:",
    }
    for chave, marcador in repetidos_regex.items():
        val = campos.get(chave, "")
        if not val:
            continue
        up = val.upper()
        matches = list(re.finditer(marcador, up))
        if len(matches) > 1:
            campos[chave] = val[: matches[1].start()].strip()
        elif len(matches) == 1 and matches[0].start() > 0:
            campos[chave] = val[: matches[0].start()].strip()

    # remove prefixos duplicados que sobram no valor
    campos["sugesp_unidade"] = corta_no_rotulo(
        remove_prefixo(campos.get("sugesp_unidade", ""), r"UNIDADE:"), r"UNIDADE:"
    )
    # quando o PDF cola "2026 SUB" na mesma linha, corta antes do ano/sub
    unidade_val = campos.get("sugesp_unidade", "")
    if unidade_val:
        m = re.search(r"\b20\d{2}\b\s*SUB", unidade_val, flags=re.IGNORECASE)
        if m:
            unidade_val = unidade_val[: m.start()].strip()
        m = re.search(r"\bSUB\s+UNIDADE:", unidade_val, flags=re.IGNORECASE)
        if m:
            unidade_val = unidade_val[: m.start()].strip()
        campos["sugesp_unidade"] = unidade_val
    campos["sugesp_sub_unidade"] = corta_no_rotulo(
        remove_prefixo(campos.get("sugesp_sub_unidade", ""), r"SUB\s+UNIDADE:"), r"SUB\s+UNIDADE:"
    )
    campos["sugesp_setor_lotacao"] = corta_no_rotulo(
        remove_prefixo(campos.get("sugesp_setor_lotacao", ""), r"SETOR DE LOTA[ÇC][ÃA]O:"),
        r"SETOR DE LOTA[ÇC][ÃA]O:",
    )
    campos["sugesp_servidor"] = corta_no_rotulo(
        remove_prefixo(campos.get("sugesp_servidor", ""), r"SERVIDOR:"), r"SERVIDOR:"
    )
    campos["sugesp_cargo"] = corta_no_rotulo(
        remove_prefixo(campos.get("sugesp_cargo", ""), r"CARGO:"), r"CARGO:"
    )
    campos["sugesp_endereco"] = corta_no_rotulo(
        remove_prefixo(campos.get("sugesp_endereco", ""), r"ENDERE[ÇC]O:"), r"ENDERE[ÇC]O:"
    )
    campos["sugesp_telefone"] = corta_no_rotulo(
        remove_prefixo(campos.get("sugesp_telefone", ""), r"TELEFONE:"), r"TELEFONE:"
    )
    campos["sugesp_email"] = corta_no_rotulo(
        remove_prefixo(campos.get("sugesp_email", ""), r"EMAIL:"), r"EMAIL:"
    )
    campos["sugesp_cpf"] = corta_no_rotulo(
        remove_prefixo(campos.get("sugesp_cpf", ""), r"CPF:"), r"CPF:"
    )

    # ano para inteiro, se possivel
    try:
        campos["sugesp_ano"] = int(campos["sugesp_ano"])
    except Exception:
        campos["sugesp_ano"] = ""

    # saneia mes
    mes_upper = campos.get("sugesp_mes_label", "")
    if mes_upper in MESES:
        campos["sugesp_mes_label"] = mes_upper
    elif mes_upper:
        for nome in MESES.keys():
            if nome.startswith(mes_upper[:3]):
                campos["sugesp_mes_label"] = nome
                break

    return campos


def _ler_upload(arquivo) -> str:
    """Lê PDF ou DOCX enviado e devolve texto contínuo."""
    if not arquivo:
        return ""
    nome = arquivo.name.lower()
    dados = arquivo.read()
    if nome.endswith(".pdf"):
        if PdfReader is None:
            st.error("Instale PyPDF2 (pip install PyPDF2) para ler PDFs.")
            return ""
        reader = PdfReader(BytesIO(dados))
        textos = []
        for pagina in reader.pages:
            textos.append(pagina.extract_text() or "")
        return "\n".join(textos)
    if nome.endswith(".docx"):
        if Document is None:
            st.error("Instale python-docx (pip install python-docx) para ler DOCX.")
            return ""
        try:
            doc = Document(BytesIO(dados))
            linhas = []
            # parágrafos soltos
            for p in doc.paragraphs:
                if p.text and p.text.strip():
                    linhas.append(p.text)
            # textos das tabelas
            for tabela in doc.tables:
                for linha in tabela.rows:
                    for celula in linha.cells:
                        if celula.text and celula.text.strip():
                            linhas.append(celula.text)
            return "\n".join(linhas)
        except Exception as e:
            st.error(f"Erro ao ler DOCX: {e}")
            return ""
    return ""


def gerar_pdf(
    ano,
    mes,
    he,
    hs,
    endereco,
    cep,
    telefone,
    data_preenchimento,
    secretaria,
    reeducando,
    funcao,
    data_inclusao,
    municipio,
    cpf,
    banco,
    agencia,
    conta,
    tipo_conta,
    feriados=None,
):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)

    # Define título interno do PDF
    c.setTitle("Folha de ponto")

    # Cabeçalho oficial
    y_top = desenhar_cabecalho(c)

    # Corpo
    desenhar_tabela(
        c,
        ano=ano,
        mes=mes,
        he=he,
        hs=hs,
        y_top=y_top,
        feriados=feriados or {},
        endereco=endereco,
        cep=cep,
        telefone=telefone,
        data_preenchimento=data_preenchimento,
        secretaria=secretaria,
        reeducando=reeducando,
        funcao=funcao,
        data_inclusao=data_inclusao,
        municipio=municipio,
        cpf=cpf,
        banco=banco,
        agencia=agencia,
        conta=conta,
        tipo_conta=tipo_conta,
    )

    c.showPage()
    c.save()

    buffer.seek(0)
    return buffer.getvalue()


def gerar_relatorio_pdf(
    ano,
    mes,
    secretaria,
    reeducando,
    funcao,
    municipio,
    endereco,
    cep,
    telefone,
    data_preenchimento,
    feriados=None,
    rodape_titulo=None,
    rodape_endereco=None,
    rodape_fone=None,
    rodape_cep=None,
    rodape_email=None,
):
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)

    # cabeçalho oficial com logo e título
    y_base = gerar_relatorio_cabecalho(
        c,
        secretaria=secretaria,
        mes=mes,
        ano=ano,
        reeducando=reeducando,
        funcao=funcao,
        municipio=municipio,
        endereco=endereco,
        cep=cep,
        telefone=telefone,
        data_preenchimento=data_preenchimento,
    )

    # tabela de atividades (dias do mês)
    atividade_base = (
        "Serviços de limpeza e conservação do prédio, bens materiais e utensílios da "
        "ULSAV/IDARON de {municipio}. Sob supervisão de um servidor."
    )
    from pdf.relatorio import desenhar_tabela_relatorio  # lazy import

    desenhar_tabela_relatorio(
        c,
        mes=mes,
        ano=ano,
        municipio=municipio,
        atividade_base=atividade_base,
        feriados=feriados or {},
        y_top=y_base,
    )

    # rodapé
    desenhar_rodape(
        c,
        titulo=rodape_titulo or "ULSAV - UNIDADE LOCAL DE SANIDADE ANIMAL E VEGETAL",
        linha_endereco=rodape_endereco or "Av. São Paulo, 436 – Bairro Centro",
        linha_fone=rodape_fone or "Fone/Fax: (69) 3642-1026/8479-9229",
        linha_cep=rodape_cep or "CEP 76.932-000 – São Miguel do Guaporé/RO",
        linha_email=rodape_email or "saomiguel@idaron.ro.gov.br",
    )

    c.showPage()
    c.save()

    buffer.seek(0)
    return buffer.getvalue()

VEICULO_MESES = [
    "Janeiro",
    "Fevereiro",
    "Marco",
    "Abril",
    "Maio",
    "Junho",
    "Julho",
    "Agosto",
    "Setembro",
    "Outubro",
    "Novembro",
    "Dezembro",
]
VEICULO_ANOS = list(range(2026, 2037))


def build_pdf_veiculo(data: dict, logo_path: Path) -> bytes:
    buffer = BytesIO()
    page_width, page_height = landscape(A4)
    c = canvas.Canvas(buffer, pagesize=(page_width, page_height))

    margin = 10 * mm
    rect_h = 30 * mm
    rect_w = page_width - 2 * margin
    rect_x = margin
    rect_y = page_height - margin - rect_h

    c.setLineWidth(0.7)
    c.rect(rect_x, rect_y, rect_w, rect_h)

    if logo_path.exists():
        logo = ImageReader(str(logo_path))
        img_w, img_h = logo.getSize()
        max_w = rect_w * 0.8
        max_h = 32 * mm
        scale = min(max_w / img_w, max_h / img_h)
        draw_w = img_w * scale
        draw_h = img_h * scale
        img_x = rect_x + (rect_w - draw_w) / 2
        img_y = rect_y + (rect_h - draw_h) / 2 - 2 * mm
        c.drawImage(logo, img_x, img_y, width=draw_w, height=draw_h, mask="auto")

    title_h = 6 * mm
    title_y = rect_y - title_h

    c.setLineWidth(0.7)
    c.rect(rect_x, title_y, rect_w, title_h)
    c.setFont("Helvetica-Bold", 10)
    c.drawCentredString(
        rect_x + rect_w / 2, title_y + 2.2 * mm, "CONTROLE DE USO E SAIDA DE VEICULO"
    )

    info_y = title_y - 4 * mm
    info_h = 8 * mm
    c.rect(rect_x, info_y - info_h, rect_w, info_h)
    c.setFont("Helvetica-Bold", 8)
    c.drawString(rect_x + 2 * mm, info_y - 7 * mm, f"ANO: {data['ano']} / MES: {data['mes']}")
    c.drawString(
        rect_x + 55 * mm,
        info_y - 7 * mm,
        f"NOME DA UNIDADE: {data['unidade'] or '______________________________'}",
    )
    c.drawString(
        rect_x + 140 * mm,
        info_y - 7 * mm,
        f"PLACA DO VEICULO: {data['placa']} ({data['modelo']})",
    )

    table_top = info_y - info_h
    bottom_area_h = 30 * mm
    table_bottom = margin + bottom_area_h
    table_height = table_top - table_bottom
    rows = 14
    row_h = table_height / rows

    col_widths = [18, 18, 18, 18, 16, 48, 82, 44]
    scale = rect_w / sum(col_widths)
    col_widths = [w * scale for w in col_widths]

    x = rect_x
    c.setLineWidth(0.5)
    for w in col_widths:
        c.line(x, table_bottom, x, table_top)
        x += w
    c.line(rect_x + rect_w, table_bottom, rect_x + rect_w, table_top)

    for i in range(rows + 1):
        y = table_top - i * row_h
        c.line(rect_x, y, rect_x + rect_w, y)

    headers = [
        "DATA",
        "HR SAIDA",
        "KM SAIDA",
        "HR CHEG",
        "KM CHEGADA",
        "DESTINO",
        "SERVICO REALIZADO DETALHADAMENTE",
        "NOME DO CONDUTOR POR EXTENSO",
    ]
    c.setFont("Helvetica-Bold", 6)
    x = rect_x
    y = table_top - row_h + 2 * mm
    for w, h in zip(col_widths, headers):
        c.drawCentredString(x + w / 2, y, h)
        x += w

    c.setFont("Helvetica", 7)
    first_col_x = rect_x + col_widths[0] / 2
    for i in range(1, rows):
        row_y = table_top - (i + 0.5) * row_h
        c.drawCentredString(first_col_x, row_y, "/    /")

    checklist_x = rect_x
    checklist_y = table_bottom - 6 * mm
    c.setFont("Helvetica-Bold", 7)
    c.drawString(checklist_x, checklist_y, "CHECKLIST:")
    c.setFont("Helvetica", 7)
    checklist_items = [
        "DOCUMENTO DE PORTE OBRIGATORIO ( )SIM ( )NAO",
        "CHAVE DE RODA ( )SIM ( )NAO",
        "MACACO ( )SIM ( )NAO",
        "TRIANGULO ( )SIM ( )NAO",
        "EXTINTOR ( )SIM ( )NAO",
        "ESTEPE ( )SIM ( )NAO",
    ]
    for idx, item in enumerate(checklist_items):
        c.drawString(checklist_x, checklist_y - (4 * mm) * (idx + 1), item)

    sig_y = margin + 8 * mm
    sig_x = rect_x + rect_w * 0.35
    sig_w = rect_w * 0.3
    c.line(sig_x, sig_y, sig_x + sig_w, sig_y)
    c.setFont("Helvetica", 7)
    c.drawCentredString(sig_x + sig_w / 2, sig_y - 4 * mm, "Assinatura do Chefe da Unidade")

    obs_x = rect_x + rect_w * 0.75
    obs_y = margin + 18 * mm
    c.setFont("Helvetica-Bold", 7)
    c.drawString(obs_x, obs_y, "OBS.:")

    c.showPage()
    c.save()
    buffer.seek(0)
    return buffer.read()


def render_veiculos():
    st.title("Entrada e Saida de Veiculos")

    st.session_state.setdefault("veiculo_show_print", False)
    st.session_state.setdefault(
        "veiculo_form_data",
        {
            "mes": VEICULO_MESES[0],
            "ano": VEICULO_ANOS[0],
            "placa": "",
            "modelo": "",
            "unidade": "",
        },
    )

    with st.form("form_veiculo"):
        mes = st.selectbox("Mes", VEICULO_MESES)
        ano = st.selectbox("Ano", VEICULO_ANOS)
        unidade = st.text_input("Nome da unidade", placeholder="ULSAV SMG")
        placa = st.text_input("Placa do veiculo", placeholder="NDI 2293")
        modelo = st.text_input("Modelo do veiculo", placeholder="TOYOTA HILUX")
        submitted = st.form_submit_button("Gerar")

    if submitted:
        st.success("Formulario enviado")
        st.session_state["veiculo_form_data"] = {
            "mes": mes,
            "ano": ano,
            "unidade": unidade,
            "placa": placa,
            "modelo": modelo,
        }
        st.session_state["veiculo_show_print"] = True

    if st.session_state["veiculo_show_print"]:
        data = st.session_state["veiculo_form_data"]

        logo_path = Path(__file__).resolve().parent / "assets" / "logo_inferior_dir.jpg"
        logo_b64 = ""
        if logo_path.exists():
            logo_b64 = base64.b64encode(logo_path.read_bytes()).decode("ascii")

        st.markdown("## Pagina de impressao")
        pdf_bytes = build_pdf_veiculo(data, logo_path)
        st.download_button(
            "Baixar PDF",
            data=pdf_bytes,
            file_name="controle_uso_saida_veiculo.pdf",
            mime="application/pdf",
        )

        st.markdown(
            f"""
            <style>
            :root {{
                --a4-landscape-width: 297mm;
                --a4-landscape-height: 210mm;
            }}
            .print-page {{
                width: min(100%, var(--a4-landscape-width));
                height: var(--a4-landscape-height);
                border: 1px solid #999;
                padding: 10mm;
                box-sizing: border-box;
                margin: 0 auto;
                background: #fff;
            }}
            .top-rect {{
                width: 100%;
                height: 30mm;
                border: 1px solid #222;
                display: flex;
                flex-direction: column;
                align-items: center;
                justify-content: center;
                box-sizing: border-box;
                margin-bottom: 0;
            }}
            .titulo-linha {{
                width: 100%;
                height: 6mm;
                display: flex;
                align-items: center;
                justify-content: center;
                font-weight: 700;
                letter-spacing: 0.5px;
                margin-bottom: 6mm;
                border-left: 1px solid #222;
                border-right: 1px solid #222;
                border-bottom: 1px solid #222;
            }}
            .top-rect img {{
                max-height: 32mm;
                max-width: 80%;
                object-fit: contain;
                margin-bottom: 2mm;
            }}
            .info-row {{
                border: 1px solid #222;
                border-top: none;
                height: 8mm;
                display: flex;
                align-items: center;
                font-size: 10px;
                padding: 0 4mm;
                gap: 12mm;
                box-sizing: border-box;
            }}
            .tabela {{
                width: 100%;
                border-collapse: collapse;
                font-size: 10px;
            }}
            .tabela th, .tabela td {{
                border: 1px solid #222;
                padding: 1mm 1mm;
                height: 7mm;
            }}
            .tabela th {{
                font-weight: 700;
                text-align: center;
            }}
            .tabela td:first-child {{
                text-align: center;
                vertical-align: middle;
            }}
            .bottom-area {{
                display: grid;
                grid-template-columns: 1fr 1fr 1fr;
                gap: 6mm;
                margin-top: 12mm;
                font-size: 10px;
            }}
            .checklist {{
                font-size: 9px;
                line-height: 1.4;
            }}
            .assinatura {{
                display: flex;
                align-items: flex-end;
                justify-content: center;
            }}
            .assinatura .linha {{
                width: 70%;
                border-top: 1px solid #222;
                text-align: center;
                padding-top: 2mm;
                font-size: 9px;
            }}
            .obs {{
                font-size: 9px;
                justify-self: end;
            }}
            @media screen {{
                .print-page {{
                    aspect-ratio: 297 / 210;
                    box-shadow: 0 8px 24px rgba(0, 0, 0, 0.12);
                }}
            }}
            </style>
            <div class="print-page">
                <div class="top-rect">
                    {"<img src='data:image/jpg;base64," + logo_b64 + "' alt='Logo' />" if logo_b64 else ""}
                </div>
                <div class="titulo-linha">CONTROLE DE USO E SAIDA DE VEICULO</div>
                <div class="info-row">
                    <div><strong>ANO:</strong> {data["ano"]} / <strong>MES:</strong> {data["mes"]}</div>
                    <div><strong>NOME DA UNIDADE:</strong> {data["unidade"] or "__________________________"}</div>
                    <div><strong>PLACA DO VEICULO:</strong> {data["placa"]} ({data["modelo"]})</div>
                </div>
                <table class="tabela">
                    <thead>
                        <tr>
                            <th>DATA</th>
                            <th>HR SAIDA</th>
                            <th>KM SAIDA</th>
                            <th>HR CHEG</th>
                            <th>KM CHEGADA</th>
                            <th>DESTINO</th>
                            <th>SERVICO REALIZADO DETALHADAMENTE</th>
                            <th>NOME DO CONDUTOR POR EXTENSO</th>
                        </tr>
                    </thead>
                    <tbody>
                        {"".join(["<tr>" + "".join(["<td>" + ("/    /" if i == 0 else "&nbsp;") + "</td>" for i in range(8)]) + "</tr>" for _ in range(14)])}
                    </tbody>
                </table>
                <div class="bottom-area">
                    <div class="checklist">
                        <strong>CHECKLIST:</strong><br/>
                        DOCUMENTO DE PORTE OBRIGATORIO ( )SIM ( )NAO<br/>
                        CHAVE DE RODA ( )SIM ( )NAO<br/>
                        MACACO ( )SIM ( )NAO<br/>
                        TRIANGULO ( )SIM ( )NAO<br/>
                        EXTINTOR ( )SIM ( )NAO<br/>
                        ESTEPE ( )SIM ( )NAO
                    </div>
                    <div class="assinatura">
                        <div class="linha">Assinatura do Chefe da Unidade</div>
                    </div>
                    <div class="obs"><strong>OBS.:</strong></div>
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )


def render_folha_ponto():
    st.title("📄 Gerador de Folha de Ponto de Reeducandos")

    with st.expander("Importar dados (PDF ou DOCX)", expanded=False):
        arquivo = st.file_uploader("Selecione o PDF ou DOCX da ?ltima folha", type=["pdf", "docx"])
        if arquivo:
            # aplica os campos apenas uma vez por arquivo para permitir edi??es depois
            if st.session_state.get("_ultimo_upload") != arquivo.name:
                st.session_state["_upload_aplicado"] = False
                st.session_state["_ultimo_upload"] = arquivo.name

            if not st.session_state.get("_upload_aplicado", False):
                texto = _ler_upload(arquivo)
                if not texto:
                    st.warning("N?o consegui ler o arquivo enviado.")
                else:
                    campos = _parse_campos(texto)
                    st.session_state.update({k: v for k, v in campos.items() if v})
                    st.session_state["_upload_aplicado"] = True
                    st.success("Campos preenchidos a partir do arquivo.")

    with st.expander("Dados do Reeducando", expanded=True):
        secretaria_input = st.text_input("Secretaria", key="secretaria")
        reeducando_input = st.text_input("Reeducando", key="reeducando")
        funcao_input = st.text_input("Função", key="funcao")
        data_inclusao_input = st.text_input("Data da inclusão", key="data_inclusao")
        municipio_input = st.text_input("Município", key="municipio")

        col_doc = st.columns(3)
        with col_doc[0]:
            cpf_input = st.text_input("CPF", key="cpf")
        with col_doc[1]:
            banco_input = st.text_input("Banco (BCO)", key="banco")
        with col_doc[2]:
            agencia_input = st.text_input("Agência (AG)", key="agencia")
        conta_input = st.text_input("Conta", key="conta")
        tipo_opcoes = ["Corrente", "Salário", "Poupança", "Outro"]
        tipo_atual = st.session_state.get("tipo_conta", "")
        if tipo_atual not in tipo_opcoes or tipo_atual == "":
            tipo_radio = "Outro"
            tipo_outro_val = tipo_atual
        else:
            tipo_radio = tipo_atual
            tipo_outro_val = ""
        col_tipo = st.columns([1, 1])
        with col_tipo[0]:
            tipo_radio_sel = st.radio("Tipo de conta", tipo_opcoes, index=tipo_opcoes.index(tipo_radio))
        with col_tipo[1]:
            tipo_outro_val = st.text_input("Outro (se aplicável)", value=tipo_outro_val, key="tipo_conta_outro")
        # valor final do tipo de conta
        if tipo_radio_sel == "Outro":
            tipo_conta_input = tipo_outro_val
        else:
            tipo_conta_input = tipo_radio_sel

        endereco_input = st.text_input("Endereço", key="endereco")
        col_contato = st.columns(3)
        with col_contato[0]:
            cep_input = st.text_input("CEP", key="cep")
        with col_contato[1]:
            telefone_input = st.text_input("Telefone", key="telefone")
        with col_contato[2]:
            data_input = st.text_input("Data", key="data_preenchimento")

    with st.expander("Rodapé (ULSAV/IDARON)", expanded=True):
        rodape_titulo_input = st.text_input("Título do rodapé", key="rodape_titulo")
        rodape_endereco_input = st.text_input("Endereço (rodapé)", key="rodape_endereco")
        rodape_fone_input = st.text_input("Fone/Fax (rodapé)", key="rodape_fone")
        rodape_cep_input = st.text_input("CEP / Cidade (rodapé)", key="rodape_cep")
        rodape_email_input = st.text_input("Email (rodapé)", key="rodape_email")

    with st.expander("Preencher dados da folha", expanded=True):
        # saneia valores de mês/ano antes de criar os widgets
        if st.session_state.get("mes_label") not in MESES:
            st.session_state["mes_label"] = list(MESES.keys())[0]
        anos_opcoes = ANOS_OPCOES
        ano_atual_ss = st.session_state.get("ano", anos_opcoes[0])
        if not isinstance(ano_atual_ss, int):
            try:
                ano_atual_ss = int(ano_atual_ss)
            except Exception:
                ano_atual_ss = anos_opcoes[0]
        if ano_atual_ss not in anos_opcoes:
            ano_atual_ss = anos_opcoes[0]
        st.session_state["ano"] = ano_atual_ss

        col_mes, col_ano = st.columns(2)
        with col_mes:
            mes_opcoes = list(MESES.keys())
            mes_label = st.selectbox("Mês", mes_opcoes, key="mes_label")
        with col_ano:
            ano = st.selectbox("Ano", anos_opcoes, key="ano")

        # opções de horário
        opcoes_he = [
            f"{h:02d}:{m:02d}"
            for h in range(5, 10)
            for m in (0, 30)
            if (h < 9 or (h == 9 and m == 0))
        ]
        opcoes_hs = [
            f"{h:02d}:{m:02d}"
            for h in range(10, 18)
            for m in (0, 30)
            if not (h == 10 and m == 0)  # começa em 10:30
        ]

        col_he, col_hs = st.columns(2)
        with col_he:
            he_default = st.session_state.get("he", opcoes_he[0])
            he_index = _safe_index(opcoes_he, he_default, 0)
            he = st.selectbox("Horário de entrada (HE)", opcoes_he, index=he_index, key="he")
        with col_hs:
            hs_default = st.session_state.get("hs", opcoes_hs[0])
            hs_index = _safe_index(opcoes_hs, hs_default, 0)
            hs = st.selectbox("Horário de saída (HS)", opcoes_hs, index=hs_index, key="hs")

        feriados_texto = st.text_area(
            "Feriados (formato: dia-descrição, separados por vírgulas)",
            value=st.session_state.get("feriados_texto", ""),
            placeholder="1-Feriado, 2-Feriado2, 3-Feriado3",
            help=(
                "Use dia-descrição separados por vírgulas. Ex.: 1-Confraternização Universal, "
                "15-Feriado inventado, 21-Dia tal (dia entre 1 e 31). "
                "Exemplo para colar: 1-Feriado, 2-Feriado2, 3-Feriado3."
            ),
        )

        st.write(
            """
        Selecione o mês e o ano, gere o PDF com o cabeçalho oficial e depois baixe o arquivo.
        """
        )

        col_btn = st.columns(2)
        with col_btn[0]:
            if st.button("Gerar Folha de Ponto"):

                feriados_dict, erros = parse_feriados_text(feriados_texto)
                if erros:
                    st.error("Revise os feriados informados: " + "; ".join(erros))
                else:
                    st.session_state["feriados_texto"] = feriados_texto
                    st.session_state["pdf"] = gerar_pdf(
                        ano=ano,
                        mes=MESES[mes_label],
                        he=he,
                        hs=hs,
                        endereco=endereco_input,
                        cep=cep_input,
                        telefone=telefone_input,
                        data_preenchimento=data_input,
                        secretaria=secretaria_input,
                        reeducando=reeducando_input,
                        funcao=funcao_input,
                        data_inclusao=data_inclusao_input,
                        municipio=municipio_input,
                        cpf=cpf_input,
                        banco=banco_input,
                        agencia=agencia_input,
                        conta=conta_input,
                        tipo_conta=tipo_conta_input,
                        feriados=feriados_dict,
                    )
                    st.success("Folha de ponto gerada com sucesso!")
        with col_btn[1]:
            if st.button("Gerar Relatório de Atividades"):
                feriados_dict, erros = parse_feriados_text(feriados_texto)
                if erros:
                    st.error("Revise os feriados informados: " + "; ".join(erros))
                else:
                    st.session_state["feriados_texto"] = feriados_texto
                    st.session_state["relatorio_pdf"] = gerar_relatorio_pdf(
                        ano=ano,
                        mes=MESES[mes_label],
                        secretaria=secretaria_input,
                        reeducando=reeducando_input,
                        funcao=funcao_input,
                        municipio=municipio_input,
                        endereco=endereco_input,
                        cep=cep_input,
                        telefone=telefone_input,
                        data_preenchimento=data_input,
                        feriados=feriados_dict,
                        rodape_titulo=rodape_titulo_input,
                        rodape_endereco=rodape_endereco_input,
                        rodape_fone=rodape_fone_input,
                        rodape_cep=rodape_cep_input,
                        rodape_email=rodape_email_input,
                    )
                    st.success("Relatório de atividades gerado com sucesso!")

    # Botões de download
    if "pdf" in st.session_state:
        st.download_button(
            "📄 Baixar Folha de Ponto",
            data=st.session_state["pdf"],
            file_name="folha.pdf",
            mime="application/pdf",
        )
    if "relatorio_pdf" in st.session_state:
        st.download_button(
            "📄 Baixar Relatório de Atividades",
            data=st.session_state["relatorio_pdf"],
            file_name="relatorio_atividades.pdf",
            mime="application/pdf",
        )
def render_folha_ponto_sugesp():
    st.title("Gerador de Folha de Ponto - SUGESP")

    with st.expander("Importar dados (PDF ou DOCX)", expanded=False):
        arquivo = st.file_uploader("Selecione o PDF ou DOCX da ultima folha", type=["pdf", "docx"], key="sugesp_upload")
        if arquivo:
            if st.session_state.get("_sugesp_ultimo_upload") != arquivo.name:
                st.session_state["_sugesp_upload_aplicado"] = False
                st.session_state["_sugesp_ultimo_upload"] = arquivo.name

            if not st.session_state.get("_sugesp_upload_aplicado", False):
                texto = _ler_upload(arquivo)
                if not texto:
                    st.warning("Nao consegui ler o arquivo enviado.")
                else:
                    campos = _parse_campos_sugesp(texto)
                    st.session_state.update({k: v for k, v in campos.items() if v})
                    st.session_state["_sugesp_upload_aplicado"] = True
                    st.success("Campos preenchidos a partir do arquivo.")

    with st.expander("Dados do servidor (SUGESP)", expanded=True):
        unidade = st.text_input("Unidade", key="sugesp_unidade")
        sub_unidade = st.text_input("Sub unidade", key="sugesp_sub_unidade")
        setor_lotacao = st.text_input("Setor de lotacao", key="sugesp_setor_lotacao")
        servidor = st.text_input("Servidor", key="sugesp_servidor")
        matricula = st.text_input("Matricula", key="sugesp_matricula")
        sigla = st.text_input("Sigla/Local", key="sugesp_sigla")
        cargo = st.text_input("Cargo", key="sugesp_cargo")

    with st.expander("Contato", expanded=True):
        endereco = st.text_input("Endereco", key="sugesp_endereco")
        cep = st.text_input("CEP", key="sugesp_cep")
        telefone = st.text_input("Telefone", key="sugesp_telefone")
        email = st.text_input("Email", key="sugesp_email")
        cpf = st.text_input("CPF", key="sugesp_cpf")
        data_preenchimento = st.text_input("Data", key="sugesp_data_preenchimento")

    with st.expander("Mes, ano e feriados", expanded=True):
        mes_opcoes = list(MESES.keys())
        if st.session_state.get("sugesp_mes_label") not in MESES:
            st.session_state["sugesp_mes_label"] = mes_opcoes[0]
        anos_opcoes = ANOS_OPCOES
        ano_atual_ss = st.session_state.get("sugesp_ano", anos_opcoes[0])
        if not isinstance(ano_atual_ss, int):
            try:
                ano_atual_ss = int(ano_atual_ss)
            except Exception:
                ano_atual_ss = anos_opcoes[0]
        if ano_atual_ss not in anos_opcoes:
            ano_atual_ss = anos_opcoes[0]
        st.session_state["sugesp_ano"] = ano_atual_ss

        col_mes, col_ano = st.columns(2)
        with col_mes:
            mes_label = st.selectbox("Mes", mes_opcoes, key="sugesp_mes_label")
        with col_ano:
            ano = st.selectbox("Ano", anos_opcoes, key="sugesp_ano")

        opcoes_he = [
            f"{h:02d}:{m:02d}"
            for h in range(5, 10)
            for m in (0, 30)
            if (h < 9 or (h == 9 and m == 0))
        ]
        opcoes_hs = [
            f"{h:02d}:{m:02d}"
            for h in range(10, 18)
            for m in (0, 30)
            if not (h == 10 and m == 0)
        ]

        col_he, col_hs = st.columns(2)
        with col_he:
            he_default = st.session_state.get("sugesp_he", "")
            he_index = _safe_index(opcoes_he, he_default, 0)
            he = st.selectbox("Horario de entrada", opcoes_he, index=he_index, key="sugesp_he")
        with col_hs:
            hs_default = st.session_state.get("sugesp_hs", "")
            hs_index = _safe_index(opcoes_hs, hs_default, 0)
            hs = st.selectbox("Horario de saida", opcoes_hs, index=hs_index, key="sugesp_hs")

        feriados_texto = st.text_area(
            "Feriados (formato: dia-descricao, separados por virgulas)",
            value=st.session_state.get("sugesp_feriados_texto", ""),
            placeholder="1-Feriado, 2-Feriado2",
        )


def render_parcelamento():
    # tenta usar locale BR se disponivel
    try:
        locale.setlocale(locale.LC_ALL, "pt_BR.UTF-8")
    except locale.Error:
        pass

    css = """
    <style>
        /* Remove negrito dos itens do radio */
        div[data-baseweb="radio"] label {
            font-weight: normal !important;
            font-family: inherit !important;
            font-size: 1rem !important;
        }
        header {
            visibility: hidden;
        }
        /* Mantem o fundo padrão do app */
        .stApp {
            background-color: transparent;
        }
        .stApp .main .block-container h1 {
            text-align: center;
        }
        /* Oculta o botao de imprimir ao imprimir */
        @media print {
            .no-print, .print-button {
                display: none !important;
            }
        }
    </style>
    """
    st.markdown(css, unsafe_allow_html=True)

    def k(name: str) -> str:
        return f"parc_{name}"

    st.session_state.setdefault(k("auto_info_submitted"), False)
    st.session_state.setdefault(k("data_requerimento"), datetime.today())
    st.session_state.setdefault(k("data_auto"), datetime.today())
    st.session_state.setdefault(k("N_auto"), "")
    st.session_state.setdefault(k("nome_completo"), "")
    st.session_state.setdefault(k("cpf"), "")
    st.session_state.setdefault(k("endereco"), "")
    st.session_state.setdefault(k("municipio"), "")
    st.session_state.setdefault(k("valor_upf"), "124,46")
    st.session_state.setdefault(k("qtd_upf_por_animal"), 2.5)
    st.session_state.setdefault(k("qtd_upf_por_parcela"), 3.0)
    st.session_state.setdefault(k("n_animais"), 0)
    st.session_state.setdefault(
        k("prazo_defesa"),
        "Sim (Desconto de 20% pra uma parcela)",
    )

    def formatar_moeda_br(valor):
        """Converte float para BRL formatado, ex: 1234.5 -> R$ 1.234,50"""
        try:
            valor = float(valor)
        except (TypeError, ValueError):
            return "R$ 0,00"

        texto = f"{valor:,.2f}"
        texto = texto.replace(",", "X").replace(".", ",").replace("X", ".")
        return f"R$ {texto}"

    col_left, col_mid, col_right = st.columns([1, 2, 1])
    with col_mid:
        st.title("Parcelar Auto de Infracao")
        tabs = st.tabs(["Preencher Requerimento", "Tabela de Descontos"])

    with tabs[0]:
        with st.expander("Dados do Auto de Infracao", expanded=True):
            st.subheader("Dados do Auto de Infracao")
            with st.form("form_auto_info"):
                col1, col2 = st.columns(2)
                with col1:
                    st.date_input(
                        "Data do requerimento",
                        key=k("data_requerimento"),
                    )
                with col2:
                    st.date_input(
                        "Data do Auto de Infracao",
                        key=k("data_auto"),
                    )

                st.text_input(
                    "Numero do Auto de Infracao:",
                    key=k("N_auto"),
                )

                continuar = st.form_submit_button("Continuar")

            if continuar:
                st.session_state[k("auto_info_submitted")] = True

        if not st.session_state[k("auto_info_submitted")]:
            st.info(
                "Preencha os dados do auto acima e pressione Enter (ou Continuar) "
                "para liberar o restante do requerimento."
            )
            st.stop()

        data_requerimento = st.session_state[k("data_requerimento")]
        data_auto = st.session_state[k("data_auto")]

        with st.expander("Dados do Autuado", expanded=True):
            st.subheader("Dados do Autuado")
            with st.form("form_requerimento"):
                nome_completo = st.text_input("Nome completo:", key=k("nome_completo"))
                cpf = st.text_input("No do CPF:", key=k("cpf"))
                endereco = st.text_input("Endereco:", key=k("endereco"))
                municipio = st.text_input("Municipio:", key=k("municipio"))

                colA, colB, colC = st.columns(3)
                with colA:
                    valor_upf = st.text_input(
                        "Valor da UPF:",
                        value=st.session_state.get(k("valor_upf"), "124,46"),
                        key=k("valor_upf"),
                    )
                with colB:
                    qtd_upf_por_animal = st.number_input(
                        "Qtd UPF por animal/Auto:",
                        min_value=0.0,
                        step=0.5,
                        format="%.2f",
                        value=float(st.session_state.get(k("qtd_upf_por_animal"), 2.5)),
                        key=k("qtd_upf_por_animal"),
                    )
                with colC:
                    qtd_upf_por_parcela = st.number_input(
                        "Qtd minima de UPF por parcela:",
                        min_value=0.0,
                        step=0.5,
                        format="%.2f",
                        value=float(st.session_state.get(k("qtd_upf_por_parcela"), 3.0)),
                        key=k("qtd_upf_por_parcela"),
                    )

                n_animais = st.number_input(
                    "Numero de Animais/Auto de Infracao:",
                    min_value=0,
                    step=1,
                    key=k("n_animais"),
                )

                prazo_defesa_escolhido = st.radio(
                    "No prazo de defesa ate 30 dias?",
                    ("Sim (Desconto de 20% pra uma parcela)", "Nao (Desconto de 10% pra uma parcela)"),
                    key=k("prazo_defesa"),
                )

                submit_form = st.form_submit_button("Aplicar / Atualizar")

            if submit_form:
                try:
                    valor_upf_float = float(valor_upf.replace(",", "."))
                except ValueError:
                    st.error("Valor da UPF invalido, usando 0.")
                    valor_upf_float = 0.0

                st.session_state[k("valor_upf_float")] = valor_upf_float

        valor_upf_float = st.session_state.get(k("valor_upf_float"), 0.0)
        total_upf = (
            st.session_state.get(k("n_animais"), 0)
            * st.session_state.get(k("qtd_upf_por_animal"), 0)
            * valor_upf_float
        )

        if total_upf > 0:
            st.metric("Valor do Auto", formatar_moeda_br(total_upf))
        else:
            st.write("Valor do Auto: R$ 0,00")

        if valor_upf_float > 0:
            min_valor_parcela = st.session_state[k("qtd_upf_por_parcela")] * valor_upf_float
        else:
            min_valor_parcela = 0

        if total_upf >= min_valor_parcela and min_valor_parcela > 0:
            num_max_parcelas = int(total_upf // min_valor_parcela)
        else:
            num_max_parcelas = 0
        num_max_parcelas = min(num_max_parcelas, 30)

        if prazo_defesa_escolhido == "Sim (Desconto de 20% pra uma parcela)":
            desconto_mensagem = "**Desconto aplicado para prazo dentro dos 30 dias**"
            coluna_desconto = "Desconto Concedido (Integral)"
        else:
            desconto_mensagem = "**Desconto aplicado para prazo fora dos 30 dias**"
            coluna_desconto = "Desconto Concedido (metade)"

        st.markdown(desconto_mensagem)

        if num_max_parcelas > 0:
            st.write(
                f"Eh possivel parcelar em ate {num_max_parcelas} vezes, respeitando "
                f"o valor minimo de R$ {min_valor_parcela:.2f} por parcela."
            )
        else:
            st.write(
                f"O valor total e menor que o minimo exigido para uma parcela: R$ {min_valor_parcela:.2f}."
            )

        parcelas_selecionadas_df = pd.DataFrame(
            columns=["Parcela", "Valor da Parcela", "Data de Vencimento"]
        ).set_index("Parcela")

        if st.session_state.get(k("prazo_defesa")) == "Sim (Desconto de 20% pra uma parcela)":
            coluna_desconto = "Desconto Concedido (Integral)"
        else:
            coluna_desconto = "Desconto Concedido (metade)"

        discount_percentage = 0
        valor_com_desconto = 0
        valor_parcela_final = 0

        if num_max_parcelas > 0:
            with st.expander("Opcoes de Parcelamento", expanded=True):
                data_dict = {
                    "Quantidade de Parcelas": list(range(1, 32)),
                    "Desconto Concedido (Integral)": [
                        20, 12, 11.5, 11, 10.5, 10, 9.5, 9, 8.5, 8,
                        7.5, 7, 6.5, 6, 5.5, 5, 4.5, 4, 3.5, 3,
                        2.5, 2, 1.75, 1.5, 1.25, 1, 0.75, 0.5, 0.25, 0, 0
                    ],
                    "Desconto Concedido (metade)": [
                        10, 6, 5.75, 5.5, 5.25, 5, 4.75, 4.5, 4.25, 4,
                        3.75, 3.5, 3.25, 3, 2.75, 2.5, 2.25, 2, 1.75, 1.5,
                        1.25, 1, 0.875, 0.75, 0.625, 0.5, 0.375, 0.25, 0.125, 0, 0
                    ],
                }

                df_descontos = pd.DataFrame(data_dict)
                df_descontos["Desconto (%)"] = df_descontos[coluna_desconto].apply(lambda x: f"{x}%")
                df_descontos["Valor com Desconto"] = total_upf * (1 - df_descontos[coluna_desconto] / 100)
                df_descontos["Valor da Parcela"] = (
                    df_descontos["Valor com Desconto"] / df_descontos["Quantidade de Parcelas"]
                )
                df_descontos["Desconto Concedido"] = total_upf - df_descontos["Valor com Desconto"]

                df_descontos_display = df_descontos.copy()
                for c in ["Valor com Desconto", "Valor da Parcela", "Desconto Concedido"]:
                    df_descontos_display[c] = df_descontos_display[c].apply(formatar_moeda_br)

                tabela_para_exibir = df_descontos_display[
                    ["Quantidade de Parcelas", "Desconto (%)", "Desconto Concedido", "Valor com Desconto", "Valor da Parcela"]
                ]
                tabela_para_exibir = tabela_para_exibir[
                    tabela_para_exibir["Quantidade de Parcelas"] <= num_max_parcelas
                ]

                st.dataframe(tabela_para_exibir, use_container_width=True)

                if num_max_parcelas > 1:
                    num_parcelas_selecionadas = st.slider(
                        "Quantidade de parcelas desejada",
                        min_value=1,
                        max_value=num_max_parcelas,
                        value=1,
                    )
                else:
                    num_parcelas_selecionadas = 1

                st.markdown(
                    "Use o controle acima para escolher o parcelamento que melhor atende o requerimento."
                )

                discount_row = df_descontos[
                    df_descontos["Quantidade de Parcelas"] == num_parcelas_selecionadas
                ].iloc[0]
                discount_percentage = discount_row[coluna_desconto]
                valor_com_desconto = total_upf * (1 - discount_percentage / 100)
                valor_parcela_final = valor_com_desconto / num_parcelas_selecionadas

                dados_parcelas = []
                for i in range(1, num_parcelas_selecionadas + 1):
                    data_venc = data_requerimento + pd.DateOffset(months=i - 1)
                    dados_parcelas.append(
                        {
                            "Parcela": i,
                            "Valor da Parcela": formatar_moeda_br(valor_parcela_final),
                            "Data de Vencimento": data_venc.strftime("%d/%m/%Y"),
                        }
                    )

                parcelas_selecionadas_df = pd.DataFrame(dados_parcelas).set_index("Parcela")

        if not parcelas_selecionadas_df.empty:
            data_req_label = st.session_state[k("data_requerimento")].strftime("%d/%m/%Y")
            data_auto_label = st.session_state[k("data_auto")].strftime("%d/%m/%Y")
            n_auto = st.session_state.get(k("N_auto"), "")
            nome_completo = st.session_state.get(k("nome_completo"), "")
            cpf = st.session_state.get(k("cpf"), "")
            endereco = st.session_state.get(k("endereco"), "")
            municipio = st.session_state.get(k("municipio"), "")

            total_upf_float = st.session_state.get(k("valor_upf_float"), 0.0)
            total_upf = (
                st.session_state.get(k("n_animais"), 0)
                * st.session_state.get(k("qtd_upf_por_animal"), 0)
                * total_upf_float
            )

            desconto_reais = total_upf - valor_com_desconto
            if desconto_reais < 0:
                desconto_reais = 0

            texto_requerimento = (
                f"Eu, {nome_completo}, brasileiro(a), portador(a) do CPF no {cpf}, "
                f"residente no endereco {endereco}, municipio de {municipio}, "
                f"venho, por meio deste requerimento datado de {data_req_label}, solicitar o parcelamento "
                f"do Auto de Infracao no {n_auto}, lavrado em {data_auto_label}, nos termos da legislacao vigente."
            )

            if total_upf > 0 and parcelas_selecionadas_df.shape[0] > 0:
                texto_parcelamento = (
                    f"O requerente solicitou o parcelamento em {parcelas_selecionadas_df.shape[0]} vezes, "
                    f"conforme a tabela de descontos, o que lhe confere o direito a um desconto de "
                    f"{discount_percentage}% (equivalente a {formatar_moeda_br(desconto_reais)}) sobre o valor inicial. "
                    f"Assim, o valor total, que originalmente era de {formatar_moeda_br(total_upf)}, "
                    f"passara a ser de {formatar_moeda_br(valor_com_desconto)}, distribuido em "
                    f"{parcelas_selecionadas_df.shape[0]} parcelas de {formatar_moeda_br(valor_parcela_final)} cada."
                )
            else:
                texto_parcelamento = (
                    "Nao e possivel parcelar, pois o valor total e inferior ao minimo exigido para uma parcela."
                )

            html = f"""
            <!DOCTYPE html>
            <html lang="pt-BR">
            <head>
                <meta charset="UTF-8">
                <title>Requerimento de Parcelamento</title>
                <style>
                    @page {{
                        margin: 20mm;
                        @bottom-center {{
                            content: "Pagina " counter(page) " de " counter(pages);
                            font-size: 10pt;
                        }}
                    }}
                    body {{
                        font-family: Arial, sans-serif;
                        margin: 20px;
                        padding: 20px;
                    }}
                    p {{
                        text-indent: 2em;
                    }}
                    .container {{
                        max-width: 800px;
                        margin: auto;
                        padding: 20px;
                        border: 1px solid #ccc;
                        border-radius: 10px;
                    }}
                    h2 {{
                        text-align: center;
                    }}
                    .texto-requerimento {{
                        margin-top: 20px;
                        line-height: 1.5;
                        text-align: justify;
                    }}
                    .texto-parcelamento {{
                        margin-top: 20px;
                        font-weight: bold;
                        text-align: justify;
                    }}
                    table {{
                        width: 100%;
                        border-collapse: collapse;
                        margin-top: 20px;
                    }}
                    th, td {{
                        border: 1px solid #ddd;
                        padding: 10px;
                        text-align: center;
                    }}
                    th {{
                        background-color: #f4f4f4;
                    }}
                    .signature {{
                        margin-top: 40px;
                        text-align: center;
                    }}
                    .signature p {{
                        margin: 0;
                        text-align: center;
                    }}
                    .print-button {{
                        display: block;
                        text-align: center;
                        margin-top: 20px;
                    }}
                    @media print {{
                        .no-print, .print-button {{
                            display: none !important;
                        }}
                    }}
                </style>
            </head>
            <body>
                <div class="container">
                    <h2>Requerimento para Parcelamento de Auto de Infracao - Emitido pela Agencia IDARON</h2>

                    <div class="texto-requerimento">
                        <p>{texto_requerimento}</p>
                    </div>

                    <div class="texto-parcelamento">
                        <p>{texto_parcelamento}</p>
                    </div>

                    <h3>Parcelas e Vencimentos</h3>
                    <table>
                        <tr>
                            <th>Parcela</th>
                            <th>Valor da Parcela</th>
                            <th>Data de Vencimento</th>
                        </tr>
            """
            for index, row in parcelas_selecionadas_df.iterrows():
                html += f"""
                        <tr>
                            <td>{index}</td>
                            <td>{row['Valor da Parcela']}</td>
                            <td>{row['Data de Vencimento']}</td>
                        </tr>
                """
            html += f"""
                    </table>

                    <div class="signature">
                        <p>Segue assinado</p>
                        <br><br>
                        <p>________________________________________</p>
                        <p>{nome_completo}</p>
                        <p>{cpf}</p>
                    </div>

                    <div class="print-button no-print">
                        <button onclick="window.print()">Imprimir Requerimento</button>
                    </div>
                </div>
            </body>
            </html>
            """

            components.html(html, height=800, scrolling=True)

    with tabs[1]:
        st.markdown("### Tabela de Descontos")
        dados = {
            "Quantidade de Parcelas": range(1, 31),
            "Desconto Concedido (Integral)": [
                20, 12, 11.5, 11, 10.5, 10, 9.5, 9, 8.5, 8,
                7.5, 7, 6.5, 6, 5.5, 5, 4.5, 4, 3.5, 3,
                2.5, 2, 1.75, 1.5, 1.25, 1, 0.75, 0.5, 0.25, 0
            ],
            "Desconto Concedido (metade)": [
                10, 6, 5.75, 5.5, 5.25, 5, 4.75, 4.5, 4.25, 4,
                3.75, 3.5, 3.25, 3, 2.75, 2.5, 2.25, 2, 1.75, 1.5,
                1.25, 1, 0.875, 0.75, 0.625, 0.5, 0.375, 0.25, 0.125, 0
            ],
        }
        df_desc = pd.DataFrame(dados)
        df_html = df_desc.to_html(index=False)
        df_html_styled = f"<style>td, th {{text-align: center;}}</style>{df_html}"
        st.markdown(df_html_styled, unsafe_allow_html=True)

    if "sugesp_pdf" in st.session_state:
        st.download_button(
            "Baixar Folha SUGESP",
            data=st.session_state["sugesp_pdf"],
            file_name="folha_sugesp.pdf",
            mime="application/pdf",
        )


col_esq, col_meio, col_dir = st.columns([1, 2, 1])
with col_meio:
    destino = option_menu(
        None,
        [
            "Folha Reeducandos",
            "Folha SUGESP",
            "Controle de Veiculos",
            "Parcelar Auto de Infracao",
        ],
        icons=["file-earmark-text", "file-earmark-ruled", "truck", "receipt"],
        menu_icon="cast",
        default_index=0,
        orientation="horizontal",
    )

if destino == "Folha Reeducandos":
    render_folha_ponto()
elif destino == "Folha SUGESP":
    render_folha_ponto_sugesp()
elif destino == "Controle de Veiculos":
    render_veiculos()
else:
    render_parcelamento()
