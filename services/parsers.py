import re
from io import BytesIO
from unicodedata import normalize

import streamlit as st

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None  # type: ignore

try:
    from docx import Document
except ImportError:
    Document = None  # type: ignore

from services.constants import MESES

def _safe_index(options, value, default=0):
    try:
        return options.index(value)
    except Exception:
        return default


def parse_feriados_text(texto):
    feriados_dict = {}
    erros = []
    for raw_bloco in texto.split(","):
        bloco = raw_bloco.strip()
        if not bloco:
            continue
        if "-" not in bloco:
            erros.append(f'"{bloco}" (faltou "-")')
            continue
        dia_str, _, nome = bloco.partition("-")
        dia_str = dia_str.strip()
        nome = nome.strip()
        try:
            dia = int(dia_str)
        except ValueError:
            erros.append(f'"{bloco}" (dia inválido)')
            continue
        if not (1 <= dia <= 31):
            erros.append(f'"{bloco}" (dia fora de 1-31)')
            continue
        if not nome:
            erros.append(f'"{bloco}" (descrição vazia)')
            continue
        feriados_dict[dia] = nome
    return feriados_dict, erros


def _clean_text(texto: str) -> str:
    # normaliza espaços e sobe para maiúsculas para regex
    return re.sub(r"\s+", " ", texto).strip()


def _parse_campos(texto: str) -> dict:
    """Extrai campos do texto plano (PDF ou DOCX)."""
    norm = _clean_text(texto).upper()
    campos = {}

    def normaliza_data(txt: str) -> str:
        if "_" in txt:
            return "__/__/____"
        bruto = txt.strip()
        digitos = re.sub(r"\D", "", bruto)
        if len(digitos) == 8:
            return f"{digitos[:2]}/{digitos[2:4]}/{digitos[4:]}"
        if not digitos:
            return "__/__/____"
        return "__/__/____"

    def normaliza_tipo_conta(txt: str) -> str:
        t = txt.upper()
        t_norm = re.sub(r"\s+", " ", t)
        # prioridade: opção marcada com (X)
        if re.search(r"\(\s*X\s*\)\s*CORRENTE", t_norm):
            return "Corrente"
        if re.search(r"\(\s*X\s*\)\s*SAL[AA]RIO", t_norm):
            return "Salário"
        if re.search(r"\(\s*X\s*\)\s*POUPAN[CC]A", t_norm):
            return "Poupança"
        # fallback por palavra-chave
        if "POUP" in t:
            return "Poupança"
        if "SAL" in t:  # SALÁRIO
            return "Salário"
        if "CORRENTE" in t:
            return "Corrente"
        return txt

    def pega(padrao, grupo=1):
        m = re.search(padrao, norm)
        return m.group(grupo).strip() if m else ""

    campos["secretaria"] = pega(r"SECRETARIA:\s*(.+?)(?:\s+ANO:|$)")
    campos["ano"] = pega(r"ANO:\s*(\d{4})")
    campos["reeducando"] = pega(r"REEDUCANDO:\s*(.+?)(?:\s+M[ÊE]S:|$)")
    campos["mes_label"] = pega(r"M[ÊE]S:\s*([A-ZÇÃÕ]+)")
    campos["funcao"] = pega(r"FUNÇÃO:\s*(.+?)(?:\s+DATA DA INCLUSÃO:|$)")
    campos["data_inclusao"] = pega(r"DATA DA INCLUS[ÃA]O:\s*([\d/]+)")
    campos["municipio"] = pega(r"MUNIC[IÍ]PIO:\s*(.+?)(?:\s+CPF:|$)")
    campos["cpf"] = pega(r"CPF:\s*([\d.\-]+)")
    campos["banco"] = pega(r"BCO:\s*([A-Z0-9]+)")
    campos["agencia"] = pega(r"AG:\s*([A-Z0-9.\-]+)")
    campos["conta"] = pega(r"CONTA:\s*([A-Z0-9.\-]+)")
    campos["tipo_conta"] = normaliza_tipo_conta(pega(r"TIPO DE CONTA:\s*(.+)"))
    campos["endereco"] = pega(r"ENDEREÇO:\s*(.+?)(?:\s+CEP:|$)")
    campos["cep"] = pega(r"CEP:\s*([\d.\-]+)")
    campos["telefone"] = pega(r"TELEFONE:\s*([0-9\s\-]+)")
    campos["data_preenchimento"] = normaliza_data(pega(r"DATA:\s*([0-9_/]+)"))

    # remove repetições quando o texto veio duplicado no DOCX/PDF
    repetidos = {
        "secretaria": "SECRETARIA:",
        "reeducando": "REEDUCANDO:",
        "funcao": "FUNÇÃO:",
        "municipio": "MUNICÍPIO:",
        "endereco": "ENDEREÇO:",
    }
    for chave, marcador in repetidos.items():
        val = campos.get(chave, "")
        if val:
            up = val.upper()
            idx = up.find(marcador, len(marcador))  # procura segunda ocorrência
            if idx > 0:
                campos[chave] = val[:idx].strip()

    # ano para inteiro, se possível
    try:
        campos["ano"] = int(campos["ano"])
    except Exception:
        campos["ano"] = ""

    # saneia mês para opção conhecida
    mes_upper = campos.get("mes_label", "")
    if mes_upper in MESES:
        campos["mes_label"] = mes_upper
    elif mes_upper:
        # tenta capitalizar conforme dicionário
        for nome in MESES.keys():
            if nome.startswith(mes_upper[:3]):
                campos["mes_label"] = nome
                break

    return campos


def _strip_accents(texto: str) -> str:
    return "".join(
        char for char in normalize("NFD", texto or "") if not re.match(r"[\u0300-\u036f]", char)
    )


def _parse_autorizacao_viagem_manual_campos(texto: str) -> dict:
    """Extrai dados do servidor de uma autorizacao de viagem manual ja emitida."""
    norm = _clean_text(texto)
    norm_upper = _strip_accents(norm).upper()
    campos = {}

    def pega(padrao: str, grupo=1) -> str:
        m = re.search(padrao, norm_upper)
        return m.group(grupo).strip() if m else ""

    campos["avm_servidor"] = pega(r"SERVIDOR:\s*(.+?)(?:\s+CARGO/FUN|$)")
    campos["avm_cargo_funcao"] = pega(r"CARGO/FUN\S*:\s*(.+?)(?:\s+MATR|$)")
    campos["avm_matricula"] = pega(r"MATR\S*CULA:\s*(.+?)(?:\s+HABILITA|$)")
    campos["avm_habilitacao"] = pega(r"HABILITA\S*:\s*(.+?)(?:\s+-\s*CATEGORIA:|$)")
    campos["avm_categoria"] = pega(r"CATEGORIA:\s*(.+?)(?:\s+-\s*VALIDADE:|$)")
    campos["avm_validade"] = pega(r"VALIDADE:\s*(.+?)(?:\s+SA|$)")
    campos["avm_responsavel_transporte"] = pega(
        r"RESPONS\S*VEL\s+TRANSPORTE:\s*(.+?)(?:\s+CHEGADA:|\s+OBS:|$)"
    )

    return campos


def _ler_upload(arquivo) -> str:
    """Lê PDF ou DOCX enviado e devolve texto contínuo."""
    if not arquivo:
        return ""
    nome = arquivo.name.lower()
    dados = arquivo.read()
    if nome.endswith(".pdf"):
        if PdfReader is None:
            st.error("Instale PyPDF2 (pip install PyPDF2) para ler PDFs.")
            return ""
        reader = PdfReader(BytesIO(dados))
        textos = []
        for pagina in reader.pages:
            textos.append(pagina.extract_text() or "")
        return "\n".join(textos)
    if nome.endswith(".docx"):
        if Document is None:
            st.error("Instale python-docx (pip install python-docx) para ler DOCX.")
            return ""
        try:
            doc = Document(BytesIO(dados))
            linhas = []
            # parágrafos soltos
            for p in doc.paragraphs:
                if p.text and p.text.strip():
                    linhas.append(p.text)
            # textos das tabelas
            for tabela in doc.tables:
                for linha in tabela.rows:
                    for celula in linha.cells:
                        if celula.text and celula.text.strip():
                            linhas.append(celula.text)
            return "\n".join(linhas)
        except Exception as e:
            st.error(f"Erro ao ler DOCX: {e}")
            return ""
    return ""
